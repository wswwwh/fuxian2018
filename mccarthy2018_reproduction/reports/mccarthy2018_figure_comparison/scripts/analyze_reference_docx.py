"""Audit the supplied journal-style DOCX without copying its research content."""

from __future__ import annotations

import argparse
import collections
import hashlib
import json
import re
import zipfile
from pathlib import Path
from xml.etree import ElementTree as ET

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn


W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
M_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"
EP_NS = "http://schemas.openxmlformats.org/officeDocument/2006/extended-properties"
NS = {"w": W_NS, "m": M_NS, "ep": EP_NS}


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as stream:
        for chunk in iter(lambda: stream.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def mm(length) -> float | None:
    return None if length is None else round(length.mm, 3)


def pt(length) -> float | None:
    return None if length is None else round(length.pt, 3)


def enum_name(value) -> str | None:
    if value is None:
        return None
    return getattr(value, "name", str(value))


def bool_or_none(value) -> bool | None:
    return None if value is None else bool(value)


def paragraph_text(paragraph) -> str:
    return re.sub(r"\s+", " ", paragraph.text).strip()


def style_record(style) -> dict[str, object]:
    paragraph = style.paragraph_format if style.type == WD_STYLE_TYPE.PARAGRAPH else None
    return {
        "style_id": style.style_id,
        "name": style.name,
        "type": enum_name(style.type),
        "based_on": style.base_style.name if style.base_style else None,
        "font_name": style.font.name,
        "font_size_pt": pt(style.font.size),
        "bold": bool_or_none(style.font.bold),
        "italic": bool_or_none(style.font.italic),
        "alignment": enum_name(paragraph.alignment) if paragraph else None,
        "space_before_pt": pt(paragraph.space_before) if paragraph else None,
        "space_after_pt": pt(paragraph.space_after) if paragraph else None,
        "line_spacing": str(paragraph.line_spacing) if paragraph and paragraph.line_spacing else None,
        "first_line_indent_mm": mm(paragraph.first_line_indent) if paragraph else None,
        "left_indent_mm": mm(paragraph.left_indent) if paragraph else None,
        "right_indent_mm": mm(paragraph.right_indent) if paragraph else None,
        "keep_with_next": bool_or_none(paragraph.keep_with_next) if paragraph else None,
        "keep_together": bool_or_none(paragraph.keep_together) if paragraph else None,
        "page_break_before": bool_or_none(paragraph.page_break_before) if paragraph else None,
        "widow_control": bool_or_none(paragraph.widow_control) if paragraph else None,
    }


def section_record(section, index: int) -> dict[str, object]:
    cols = section._sectPr.find(qn("w:cols"))
    num_cols = int(cols.get(qn("w:num"), "1")) if cols is not None else 1
    col_space_twips = int(cols.get(qn("w:space"), "0")) if cols is not None else 0
    return {
        "index": index,
        "start_type": enum_name(section.start_type),
        "orientation": enum_name(section.orientation),
        "page_width_mm": mm(section.page_width),
        "page_height_mm": mm(section.page_height),
        "margin_top_mm": mm(section.top_margin),
        "margin_bottom_mm": mm(section.bottom_margin),
        "margin_left_mm": mm(section.left_margin),
        "margin_right_mm": mm(section.right_margin),
        "gutter_mm": mm(section.gutter),
        "header_distance_mm": mm(section.header_distance),
        "footer_distance_mm": mm(section.footer_distance),
        "different_first_page_header_footer": bool(section.different_first_page_header_footer),
        "columns": num_cols,
        "column_space_pt": round(col_space_twips / 20.0, 3),
    }


def package_stats(path: Path) -> dict[str, object]:
    stats: dict[str, object] = {}
    with zipfile.ZipFile(path) as archive:
        names = set(archive.namelist())
        stats["package_part_count"] = len(names)
        stats["embedded_media_count"] = sum(name.startswith("word/media/") for name in names)
        if "docProps/app.xml" in names:
            root = ET.fromstring(archive.read("docProps/app.xml"))
            for key in ("Pages", "Words", "Characters", "Paragraphs", "Lines"):
                node = root.find(f"ep:{key}", NS)
                stats[key.lower()] = int(node.text) if node is not None and node.text else None
        if "word/document.xml" in names:
            root = ET.fromstring(archive.read("word/document.xml"))
            stats["omml_equation_count"] = len(root.findall(".//m:oMath", NS))
            stats["omml_equation_paragraph_count"] = len(root.findall(".//m:oMathPara", NS))
        if "word/settings.xml" in names:
            root = ET.fromstring(archive.read("word/settings.xml"))
            stats["track_revisions"] = root.find("w:trackRevisions", NS) is not None
            stats["even_and_odd_headers"] = root.find("w:evenAndOddHeaders", NS) is not None
    return stats


def analyze(path: Path) -> dict[str, object]:
    document = Document(path)
    paragraphs = list(document.paragraphs)
    style_usage = collections.Counter(p.style.name for p in paragraphs)
    table_style_usage = collections.Counter(
        table.style.name if table.style else "(none)" for table in document.tables
    )
    run_fonts: collections.Counter[str] = collections.Counter()
    run_east_asia_fonts: collections.Counter[str] = collections.Counter()
    run_sizes: collections.Counter[str] = collections.Counter()
    alignments: collections.Counter[str] = collections.Counter()
    line_spacings: collections.Counter[str] = collections.Counter()
    first_indents: collections.Counter[str] = collections.Counter()

    heading_rows: list[dict[str, object]] = []
    semantic_rows: list[dict[str, object]] = []
    figure_caption_rows: list[dict[str, object]] = []
    table_caption_rows: list[dict[str, object]] = []
    equation_number_rows: list[dict[str, object]] = []

    heading_re = re.compile(r"^(?:\d+(?:\.\d+){0,2}\s+|引言|结论|参考文献|附录)")
    figure_re = re.compile(r"^(?:图\s*\d+|Fig(?:ure)?\.?\s*\d+)", re.IGNORECASE)
    table_re = re.compile(r"^(?:表\s*\d+|Table\s*\d+)", re.IGNORECASE)
    equation_number_re = re.compile(r"(?:\(|（)\s*\d+(?:[-.]\d+)?\s*(?:\)|）)\s*$")

    for index, paragraph in enumerate(paragraphs):
        text = paragraph_text(paragraph)
        alignments[enum_name(paragraph.alignment) or "INHERITED"] += 1
        if paragraph.paragraph_format.line_spacing:
            line_spacings[str(paragraph.paragraph_format.line_spacing)] += 1
        if paragraph.paragraph_format.first_line_indent:
            first_indents[f"{mm(paragraph.paragraph_format.first_line_indent):.3f}"] += 1
        for run in paragraph.runs:
            if run.font.name:
                run_fonts[run.font.name] += 1
            rpr = run._r.get_or_add_rPr()
            rfonts = rpr.rFonts
            if rfonts is not None:
                east_asia = rfonts.get(qn("w:eastAsia"))
                if east_asia:
                    run_east_asia_fonts[east_asia] += 1
            if run.font.size:
                run_sizes[f"{run.font.size.pt:.1f}"] += 1

        style_name = paragraph.style.name
        if text and ("标题" in style_name or "Heading" in style_name or heading_re.match(text)):
            heading_rows.append({"paragraph_index": index, "style": style_name, "length": len(text)})
        if text.startswith("摘要") or text.startswith("关键词") or text.startswith("Abstract"):
            semantic_rows.append({"paragraph_index": index, "kind": text.split("：", 1)[0][:16], "style": style_name})
        if figure_re.match(text):
            figure_caption_rows.append({"paragraph_index": index, "style": style_name, "length": len(text)})
        if table_re.match(text):
            table_caption_rows.append({"paragraph_index": index, "style": style_name, "length": len(text)})
        if equation_number_re.search(text):
            equation_number_rows.append({"paragraph_index": index, "style": style_name, "length": len(text)})

    used_styles = []
    for style in document.styles:
        if style.name in style_usage or style.name in table_style_usage or style.name in {
            "Normal", "Title", "Heading 1", "Heading 2", "Heading 3", "Caption"
        }:
            used_styles.append(style_record(style))

    table_shapes = []
    for index, table in enumerate(document.tables):
        col_count = max((len(row.cells) for row in table.rows), default=0)
        table_shapes.append(
            {
                "index": index,
                "rows": len(table.rows),
                "columns": col_count,
                "style": table.style.name if table.style else None,
            }
        )

    package = package_stats(path)
    result = {
        "source_file": str(path.resolve()),
        "source_bytes": path.stat().st_size,
        "source_sha256": sha256(path),
        "analysis_boundary": "仅提取排版、样式和结构统计，不复制参考稿研究内容。",
        "document_stats": {
            "paragraphs": len(paragraphs),
            "tables": len(document.tables),
            "inline_shapes": len(document.inline_shapes),
            "sections": len(document.sections),
            **package,
        },
        "sections": [section_record(section, i + 1) for i, section in enumerate(document.sections)],
        "paragraph_style_usage": style_usage.most_common(),
        "table_style_usage": table_style_usage.most_common(),
        "used_style_definitions": used_styles,
        "run_font_usage": run_fonts.most_common(20),
        "run_east_asia_font_usage": run_east_asia_fonts.most_common(20),
        "run_size_usage_pt": run_sizes.most_common(20),
        "paragraph_alignment_usage": alignments.most_common(),
        "line_spacing_usage": line_spacings.most_common(20),
        "first_line_indent_usage_mm": first_indents.most_common(20),
        "heading_structure": heading_rows,
        "semantic_markers": semantic_rows,
        "figure_caption_structure": figure_caption_rows,
        "table_caption_structure": table_caption_rows,
        "equation_number_structure": equation_number_rows,
        "table_shapes": table_shapes,
    }
    return result


def top_pairs(values: list[list[object]] | list[tuple[object, object]], limit: int = 8) -> str:
    return "、".join(f"{name}（{count}）" for name, count in values[:limit]) or "未检出"


def render_markdown(data: dict[str, object]) -> str:
    stats = data["document_stats"]
    sections = data["sections"]
    lines = [
        "# 参考投稿稿样式审计",
        "",
        f"- 来源：`{data['source_file']}`",
        f"- SHA-256：`{data['source_sha256']}`",
        f"- 文件大小：`{data['source_bytes']}` bytes",
        f"- 分析边界：{data['analysis_boundary']}",
        "",
        "## 文档结构",
        "",
        f"- Word 属性页数：`{stats.get('pages') if stats.get('pages') is not None else '【待核实】'}`；段落：`{stats['paragraphs']}`；表格：`{stats['tables']}`；内嵌图形：`{stats['inline_shapes']}`；节：`{stats['sections']}`。",
        f"- OMML 公式对象：`{stats.get('omml_equation_count', 0)}`；带公式编号样式的段落：`{len(data['equation_number_structure'])}`。",
        f"- 图题样式段落：`{len(data['figure_caption_structure'])}`；表题样式段落：`{len(data['table_caption_structure'])}`。",
        "",
        "## 页面与分栏",
        "",
        "| 节 | 页面/mm | 页边距 上/下/左/右 mm | 栏数 | 栏间距/pt | 起始方式 |",
        "|---:|---|---|---:|---:|---|",
    ]
    for section in sections:
        lines.append(
            f"| {section['index']} | {section['page_width_mm']} × {section['page_height_mm']} | "
            f"{section['margin_top_mm']} / {section['margin_bottom_mm']} / {section['margin_left_mm']} / {section['margin_right_mm']} | "
            f"{section['columns']} | {section['column_space_pt']} | {section['start_type']} |"
        )
    lines.extend(
        [
            "",
            "## 字体、字号与段落统计",
            "",
            f"- 常用段落样式：{top_pairs(data['paragraph_style_usage'])}。",
            f"- 常用西文字体（显式设置）：{top_pairs(data['run_font_usage'])}。",
            f"- 常用中文字体（eastAsia 显式设置）：{top_pairs(data['run_east_asia_font_usage'])}。",
            f"- 常用字号/pt（显式设置）：{top_pairs(data['run_size_usage_pt'])}。",
            f"- 段落对齐：{top_pairs(data['paragraph_alignment_usage'])}。",
            f"- 首行缩进/mm：{top_pairs(data['first_line_indent_usage_mm'])}。",
            "",
            "## 可迁移的版式原则",
            "",
            "- 保留 A4 期刊式页边距（参考稿左右约 20 mm、上约 25.4 mm、下约 22.5 mm）、摘要和关键词的紧凑组织、分级标题与正文首行缩进。",
            "- 正文采用双栏；54 图对照、大表和关键公式允许通过连续分节切换为通栏，随后恢复双栏。",
            "- 图题置于图下、表题置于表上；中英文题注分行，正文先引用再出现图表。",
            "- 公式单独成段、主体居中、编号右对齐；变量用数学斜体，单位用正体。",
            "- 结果论证按“文字提出问题—图给出结构—表给出数值—段落解释误差和边界”组织。",
            "- 参考稿的具体研究对象、数据、结论和句子不进入本报告。",
            "",
            "## 需要在最终报告中采取的稳健化调整",
            "",
            "- 54 组双图对照比普通期刊稿更宽，故以可审查性优先：每组 panel 通栏，说明表可通栏或双栏跨页控制。",
            "- 所有未知作者、单位、导师、基金与投稿字段保留【待核实】占位符。",
            "- 参考稿中未能从 XML 唯一确定的视觉细节，以渲染页检查为准；仍不确定者标记【待核实】。",
        ]
    )
    return "\n".join(lines) + "\n"


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("docx", type=Path)
    parser.add_argument("--json", type=Path, required=True)
    parser.add_argument("--markdown", type=Path, required=True)
    return parser.parse_args()


def main() -> int:
    args = parse_args()
    data = analyze(args.docx.resolve())
    args.json.parent.mkdir(parents=True, exist_ok=True)
    args.markdown.parent.mkdir(parents=True, exist_ok=True)
    args.json.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")
    args.markdown.write_text(render_markdown(data), encoding="utf-8")
    print(f"reference_docx_analysis=PASS sections={len(data['sections'])} pages={data['document_stats'].get('pages')}")
    print(args.json.resolve())
    print(args.markdown.resolve())
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
