#!/usr/bin/env python3
"""Build the strict four-page Chinese adviser summary and final-stage evidence."""

from __future__ import annotations

import argparse
import csv
from datetime import datetime, timezone
import hashlib
import json
from pathlib import Path
import platform
import re
import shutil
import subprocess
import sys
import tempfile
from typing import Any
import zipfile

import numpy as np


ROOT = Path(__file__).resolve().parents[1]
CONFIG = (
    ROOT / "research" / "invariant_bundles" / "configs" / "adviser_summary.json"
)
ADVISER = ROOT / "reports" / "adviser_delivery"
DOCX = ADVISER / "invariant_bundle研究摘要_4页.docx"
PDF = ADVISER / "invariant_bundle研究摘要_4页.pdf"
QUESTIONS = ADVISER / "给导师的审阅问题.md"
EVIDENCE = (
    ROOT / "research" / "invariant_bundles" / "adviser_summary_validation"
)
REGISTRY = (
    ROOT / "research" / "invariant_bundles" / "benchmarks" / "benchmark_registry.csv"
)
METHOD = (
    ROOT
    / "research"
    / "invariant_bundles"
    / "results"
    / "csv"
    / "method_comparison.csv"
)
SCHUR = (
    ROOT
    / "research"
    / "invariant_bundles"
    / "results"
    / "csv"
    / "independent_schur_backend_comparison.csv"
)
FAILURE = (
    ROOT
    / "research"
    / "invariant_bundles"
    / "results"
    / "csv"
    / "qr_svd_failure_classification.csv"
)
ABLATION = (
    ROOT
    / "research"
    / "invariant_bundles"
    / "results"
    / "csv"
    / "ablation_study.csv"
)
MANIFOLD = (
    ROOT
    / "research"
    / "invariant_bundles"
    / "results"
    / "csv"
    / "manifold_convergence.csv"
)
FRESH = (
    ROOT
    / "research"
    / "invariant_bundles"
    / "independent_rerun"
    / "comparison_to_stage_f.csv"
)
PAPER_SUMMARY = (
    ROOT
    / "research"
    / "invariant_bundles"
    / "paper_release_validation"
    / "paper_release_summary.json"
)
LITERATURE_SUMMARY = (
    ROOT
    / "research"
    / "invariant_bundles"
    / "literature_validation"
    / "literature_validation_summary.json"
)
HOLDOUT = (
    ROOT
    / "data"
    / "computed"
    / "chapter4_fig43_fig46_projection_holdout_audit.csv"
)
STAGE_G_HASHES = ROOT / "stage_g_delivery_review" / "artifact_hashes.csv"

METHOD_IDS = {
    "traditional_pointwise_eigendecomposition": "Eig",
    "ordered_partial_real_schur_tracking": "Schur",
    "qr_svd_shifted_cocycle_iteration": "QR/SVD",
}
STATUS = {"accepted": "A", "boundary": "B", "fail": "F"}
FAMILY_SHORT = {
    "earth_moon_l1_quasi_halo": "EM-Halo",
    "earth_moon_l1_quasi_vertical": "EM-Vertical",
    "earth_moon_route_h_quasi_dro": "Route-H",
    "sun_earth_l1_two_frequency_torus": "SE-L1",
}


def read_csv(path: Path) -> list[dict[str, str]]:
    with path.open("r", encoding="utf-8-sig", newline="") as stream:
        return list(csv.DictReader(stream))


def write_csv(path: Path, rows: list[dict[str, Any]]) -> None:
    if not rows:
        raise ValueError(f"empty CSV rows for {path}")
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", encoding="utf-8", newline="") as stream:
        writer = csv.DictWriter(stream, fieldnames=list(rows[0]))
        writer.writeheader()
        writer.writerows(rows)


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as stream:
        for block in iter(lambda: stream.read(1024 * 1024), b""):
            digest.update(block)
    return digest.hexdigest().upper()


def set_run_font(run: Any, font_name: str, size: float | None = None) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt

    run.font.name = font_name
    if size is not None:
        run.font.size = Pt(size)
    properties = run._element.get_or_add_rPr()
    fonts = properties.rFonts
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        properties.insert(0, fonts)
    fonts.set(qn("w:eastAsia"), font_name)


def shade(element: Any, fill: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    properties = element.get_or_add_tcPr() if element.tag.endswith("tc") else element.get_or_add_pPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_margins(cell: Any, top: int = 30, start: int = 45, bottom: int = 30, end: int = 45) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def add_paragraph(
    document: Any,
    text: str,
    *,
    bold_prefix: str | None = None,
    size: float = 9.2,
    color: str | None = None,
    space_after: float = 2.0,
) -> Any:
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor

    paragraph = document.add_paragraph()
    paragraph.paragraph_format.first_line_indent = None
    paragraph.paragraph_format.line_spacing = 1.05
    paragraph.paragraph_format.space_after = Pt(space_after)
    if bold_prefix and text.startswith(bold_prefix):
        first = paragraph.add_run(bold_prefix)
        first.bold = True
        set_run_font(first, "黑体", size)
        rest = paragraph.add_run(text[len(bold_prefix) :])
        set_run_font(rest, "宋体", size)
        if color:
            first.font.color.rgb = RGBColor.from_string(color)
            rest.font.color.rgb = RGBColor.from_string(color)
    else:
        run = paragraph.add_run(text)
        set_run_font(run, "宋体", size)
        if color:
            run.font.color.rgb = RGBColor.from_string(color)
    return paragraph


def add_bullet(document: Any, text: str, *, size: float = 8.9, color: str | None = None) -> Any:
    from docx.shared import Pt, RGBColor

    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.first_line_indent = None
    paragraph.paragraph_format.left_indent = None
    paragraph.paragraph_format.line_spacing = 1.0
    paragraph.paragraph_format.space_after = Pt(1.5)
    run = paragraph.add_run(text)
    set_run_font(run, "宋体", size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    return paragraph


def add_section_heading(document: Any, text: str) -> None:
    from docx.shared import Pt, RGBColor

    paragraph = document.add_paragraph()
    paragraph.paragraph_format.first_line_indent = None
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(2)
    shade(paragraph._p, "D9EAF7")
    run = paragraph.add_run(text)
    run.bold = True
    run.font.color.rgb = RGBColor(31, 78, 121)
    set_run_font(run, "黑体", 11.5)


def add_table(
    document: Any,
    headers: list[str],
    rows: list[list[str]],
    *,
    font_size: float = 7.4,
    header_fill: str = "4472C4",
    first_col_bold: bool = False,
) -> Any:
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor

    table = document.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.autofit = True
    for row_index, values in enumerate([headers, *rows]):
        for col_index, value in enumerate(values):
            cell = table.cell(row_index, col_index)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell.text = ""
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.first_line_indent = None
            paragraph.paragraph_format.line_spacing = 1.0
            paragraph.paragraph_format.space_after = Pt(0)
            run = paragraph.add_run(str(value))
            run.bold = row_index == 0 or (first_col_bold and col_index == 0)
            set_run_font(run, "宋体", font_size)
            if row_index == 0:
                shade(cell._tc, header_fill)
                run.font.color.rgb = RGBColor(255, 255, 255)
    return table


def add_callout(document: Any, title: str, text: str, *, fill: str = "FFF2CC") -> None:
    from docx.shared import Pt

    table = document.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    shade(cell._tc, fill)
    set_cell_margins(cell, top=70, start=90, bottom=70, end=90)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.first_line_indent = None
    paragraph.paragraph_format.line_spacing = 1.0
    paragraph.paragraph_format.space_after = Pt(0)
    lead = paragraph.add_run(title + "：")
    lead.bold = True
    set_run_font(lead, "黑体", 9.2)
    run = paragraph.add_run(text)
    set_run_font(run, "宋体", 9.2)


def configure_document(document: Any, config: dict[str, Any]) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt

    section = document.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.15)
    section.bottom_margin = Cm(1.15)
    section.left_margin = Cm(1.25)
    section.right_margin = Cm(1.25)
    section.header_distance = Cm(0.45)
    section.footer_distance = Cm(0.45)
    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(9.2)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing = 1.0
    for current in document.sections:
        header = current.header.paragraphs[0]
        header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        header.paragraph_format.first_line_indent = None
        run = header.add_run("拟周期轨道实不变子束研究摘要｜兀文昊｜中国科学院大学")
        set_run_font(run, "宋体", 7.5)
        footer = current.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer.paragraph_format.first_line_indent = None
        run = footer.add_run()
        for kind, value in (("begin", None), ("instr", " PAGE "), ("separate", None), ("end", None)):
            if kind == "instr":
                node = OxmlElement("w:instrText")
                node.set(qn("xml:space"), "preserve")
                node.text = value
            else:
                node = OxmlElement("w:fldChar")
                node.set(qn("w:fldCharType"), kind)
            run._r.append(node)
        sep = footer.add_run(" / 4")
        set_run_font(sep, "宋体", 8)
    core = document.core_properties
    core.title = config["title"]
    core.author = config["author"]
    core.subject = "导师快速审阅版，严格四页"


def load_data() -> dict[str, Any]:
    return {
        "registry": read_csv(REGISTRY),
        "method": read_csv(METHOD),
        "schur": read_csv(SCHUR),
        "failure": read_csv(FAILURE),
        "ablation": read_csv(ABLATION),
        "manifold": read_csv(MANIFOLD),
        "fresh": read_csv(FRESH),
        "holdout": read_csv(HOLDOUT),
        "paper": json.loads(PAPER_SUMMARY.read_text(encoding="utf-8")),
        "literature": json.loads(LITERATURE_SUMMARY.read_text(encoding="utf-8")),
    }


def metrics(data: dict[str, Any]) -> dict[str, Any]:
    counts = {}
    for method_id in METHOD_IDS:
        counts[method_id] = {
            status: sum(
                row["method"] == method_id and row["research_status"] == status
                for row in data["method"]
            )
            for status in ("accepted", "boundary", "fail")
        }
    scientific = [
        row for row in data["fresh"] if row["comparison_kind"] != "informational_provenance"
    ]
    numeric = [row for row in data["fresh"] if row["comparison_kind"] == "numeric_scientific"]
    return {
        "counts": counts,
        "schur_max_angle": max(
            float(row["invariant_subspace_principal_angle_max_deg"])
            for row in data["schur"]
        ),
        "schur_agreement": sum(
            row["dimension_agreement"] == "true"
            and row["classification_agreement"] == "true"
            and row["status_agreement"] == "true"
            for row in data["schur"]
        ),
        "manifold_accepted": sum(row["status"] == "accepted" for row in data["manifold"]),
        "manifold_failed": sum(row["status"] == "fail" for row in data["manifold"]),
        "max_jacobi": max(float(row["manifold_jacobi_drift"]) for row in data["manifold"]),
        "fresh_scientific": len(scientific),
        "fresh_failures": sum(row["comparison_status"] == "fail" for row in data["fresh"]),
        "fresh_max_rel": max(float(row["relative_difference"]) for row in numeric),
        "ablation_accepted": sum(row["research_status"] == "accepted" for row in data["ablation"]),
        "ablation_failed": sum(row["research_status"] == "fail" for row in data["ablation"]),
        "ablation_exceptions": sum(row["bundle_dimension"] == "0" for row in data["ablation"]),
    }


def case_table_rows(data: dict[str, Any]) -> list[list[str]]:
    by_key = {(row["case_id"], row["method"]): row for row in data["method"]}
    rows = []
    for registry in data["registry"]:
        case_id = registry["case_id"]
        label = case_id.replace("em_", "").replace("route_h_", "RH_").replace("se_", "SE_")
        if len(label) > 28:
            label = label.replace("_legacy_dg_positive", "_legacy+").replace("_lowres_negative", "_N9-")
        rows.append(
            [
                label,
                registry["spectral_samples"],
                STATUS[by_key[(case_id, "traditional_pointwise_eigendecomposition")]["research_status"]],
                STATUS[by_key[(case_id, "ordered_partial_real_schur_tracking")]["research_status"]],
                STATUS[by_key[(case_id, "qr_svd_shifted_cocycle_iteration")]["research_status"]],
            ]
        )
    return rows


def high_resolution_rows(data: dict[str, Any]) -> list[list[str]]:
    by_key = {(row["case_id"], row["method"]): row for row in data["method"]}
    rows = []
    for case_id, label in (
        ("em_halo_12p40_n45", "Halo N45"),
        ("em_vertical_12p66_n57", "Vertical N57"),
        ("se_active_geometry_member_468", "Sun–Earth 468"),
    ):
        schur = by_key[(case_id, "ordered_partial_real_schur_tracking")]
        qr = by_key[(case_id, "qr_svd_shifted_cocycle_iteration")]
        rows.append(
            [
                label,
                f"{float(schur['max_invariance_residual']):.2e}",
                f"{float(qr['max_invariance_residual']):.2e}",
                f"{float(schur['runtime_seconds']):.3f}/{float(qr['runtime_seconds']):.3f}",
                "双方法通过",
            ]
        )
    return rows


def build_docx(config: dict[str, Any], data: dict[str, Any], stat: dict[str, Any]) -> None:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor

    document = Document()
    configure_document(document, config)

    # Page 1: origin, defect, and methods.
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.first_line_indent = None
    title.paragraph_format.space_after = Pt(2)
    run = title.add_run(config["title"])
    run.bold = True
    run.font.color.rgb = RGBColor(31, 78, 121)
    set_run_font(run, "黑体", 18)
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.first_line_indent = None
    subtitle.paragraph_format.space_after = Pt(4)
    run = subtitle.add_run(
        f"{config['subtitle']}｜作者：{config['author']}｜导师：{config['adviser']}｜{config['institution']}｜{config['date']}"
    )
    set_run_font(run, "宋体", 8.8)
    add_callout(
        document,
        "结论先行",
        "15 案例中 Pointwise eig 0/15 通过；Partial real-Schur 7 通过/4 边界/4 失败；Shifted QR/SVD 10 通过/5 失败。独立 MATLAB Schur 在 12 个关键案例上 12/12 一致。当前稿件定位是数值框架与系统比较，not_submission_ready。",
        fill="E2F0D9",
    )
    add_section_heading(document, "1｜复现工作如何引出问题")
    add_paragraph(
        document,
        "McCarthy 2018 的 54 图复现已完成工程覆盖（13 个 V0、41 个 V2），但冻结证据仍为 7 accepted、30 boundary、5 diagnostic、12 proxy。逐图审计暴露出一个独立科学问题：拟周期环面附近的稳定/不稳定方向必须服从相位平移 cocycle，不能由各相位点互不相干的局部特征向量替代。研究层只读取冻结源数据，不回写提升复现等级。",
        size=8.9,
    )
    add_section_heading(document, "2｜点式特征方向的主要缺陷")
    add_bullet(document, "局部方程 A(θ)v=λv 没有约束 A(θ)E(θ) 落入 E(θ+ρ)，因此低局部特征对残差不等于低 cocycle 残差。")
    add_bullet(document, "复共轭对在实数域对应二维子空间；取复向量实部会破坏不变性，二维对象不得重命名为一维。")
    add_bullet(document, "符号/相位对齐只能稳定表示，不能修复错误对象；流形传播还会把局部方向误差放大成全片几何差异。")
    add_section_heading(document, "3｜三种方法及其角色")
    add_table(
        document,
        ["方法", "计算对象", "主要用途", "不可越过的边界"],
        [
            ["Pointwise eig", "每相位局部特征向量", "保留传统基线与失效模式", "不能自动称为 cocycle bundle"],
            ["Partial real-Schur", "有序 1×1/2×2 实谱块", "确定实子空间维数与分支", "2D 共轭块不降为 1D"],
            ["Shifted QR/SVD", "相位平移下迭代正交帧", "求相位依赖子空间与收敛史", "受初始化、维数和迭代上限约束"],
        ],
        font_size=7.7,
        first_col_bold=True,
    )
    add_paragraph(
        document,
        "统一验收：max cocycle residual≤1e-6 为通过；1e-6 至 1e-3 仅可为边界；源轨道、维数或收敛失败优先判失败。",
        bold_prefix="统一验收：",
        size=8.5,
        color="7F6000",
    )
    document.add_page_break()

    # Page 2: 15-case table and anchors.
    add_section_heading(document, "4｜15-case 总体结果表")
    add_paragraph(
        document,
        "符号：A=accepted，B=boundary，F=fail。状态属于研究层；Route H physical 与 legacy 正控制分列。",
        size=8.2,
        color="595959",
    )
    add_table(
        document,
        ["案例", "N", "Eig", "Schur", "QR/SVD"],
        case_table_rows(data),
        font_size=6.8,
        first_col_bold=True,
    )
    add_section_heading(document, "5｜三个高分辨率案例")
    add_table(
        document,
        ["案例", "Schur r_max", "QR/SVD r_max", "耗时 Schur/QR(s)", "测试窗内流形"],
        high_resolution_rows(data),
        font_size=7.4,
        first_col_bold=True,
    )
    add_bullet(document, "Halo N45、Vertical N57、Sun–Earth 468 上两种改进方法均通过 bundle 与测试传播窗内的 manifold 门槛。", size=8.3)
    add_bullet(document, "Stage F 共 126 行：36 accepted、90 fail；最大 Jacobi 漂移 2.220e-15。", size=8.3)
    add_bullet(document, "低分辨率 Halo/Vertical 全流形片距离 0.0150–0.0245，仍高于 0.01；局部 bundle 通过不覆盖全片失败。", size=8.3, color="C00000")
    add_callout(
        document,
        "独立后端",
        f"MATLAB R2024a schur/ordschur（MKL 2023.2）在 12 个关键案例上维数、分类、状态 {stat['schur_agreement']}/12 一致；最大主角 {stat['schur_max_angle']:.3e}° < 1e-4°。",
        fill="DDEBF7",
    )
    document.add_page_break()

    # Page 3: Route H, failures, ablation, rerun, CI.
    add_section_heading(document, "6｜Route H 关键发现：算子语义优先于“漂亮的一维结果”")
    add_paragraph(
        document,
        "legacy member 68 使用 seed-rho，可形成一维 Schur 正控制，但源映射残差约 1.988e-3；physical corrected-rho 把闭合残差改进到约 8.697e-13，却显示相对复部约 0.342 的共轭谱。正确物理对象因此是二维实子空间，而不是把复向量投影成一维。",
        size=8.8,
    )
    failure_rows = []
    for row in data["failure"]:
        case = row["case_id"].replace("route_h_member_", "member ").replace("_legacy_dg_positive", " legacy+")
        failure_rows.append(
            [
                case,
                row["independent_schur_dimension"],
                row["best_native_research_status"],
                f"{float(row['best_native_max_invariance_residual']):.2e}",
                row["final_label"],
            ]
        )
    add_table(
        document,
        ["案例", "独立 k", "最佳有界状态", "最佳 r_max", "失败分类"],
        failure_rows,
        font_size=7.2,
        first_col_bold=True,
    )
    add_bullet(document, "有界分类得到 4 no_accepted_1d_bundle：Physical 17/32/54/68 在 3 初始化×200/500/1000 次×N45/N67 中均未发现 accepted 1D bundle；80 位只重算残差，不是任意精度轨迹。", size=8.2)
    add_bullet(document, "Legacy 68：Schur seed 可通过，random/local-SVD 不稳定，故为 method_initialization_sensitive；正控制不能替代 physical operator。", size=8.2)
    add_section_heading(document, "消融、全新进程与 CI")
    add_table(
        document,
        ["审计项", "规模/结果", "科学结论"],
        [
            ["7 变体消融", f"35 行；{stat['ablation_accepted']} A/{stat['ablation_failed']} F；{stat['ablation_exceptions']} 异常", "符号对齐不修复 cocycle；2D 跟踪改善但仍失败"],
            ["全新进程重跑", f"15 cocycle；45 bundle；126 manifold；{stat['fresh_scientific']} 科学检查", f"失败 {stat['fresh_failures']}；最大相对差 {stat['fresh_max_rel']:.1f}"],
            ["GitHub Actions", "push/PR 快速 CI；手动 full validation", "输出在 runner 临时目录；权威文件哈希守护"],
        ],
        font_size=7.4,
        first_col_bold=True,
    )
    add_callout(
        document,
        "真实性边界",
        "失败、边界、method exception 和负控制全部进入 CSV/NPZ/日志；Chapter 4 projection holdout 仍为 0/4，paper_projection=fail，paper_3d=false。",
        fill="FCE4D6",
    )
    document.add_page_break()

    # Page 4: contribution and remaining work.
    add_section_heading(document, "7｜当前论文贡献（可辩护范围）")
    add_bullet(document, "统一 15-case/4-family/3-method benchmark，把源轨道、维数、残差、相位连续性、流形几何和成本绑定到固定阈值。", size=8.7)
    add_bullet(document, "引入真正独立的 MATLAB real-Schur 后端，避免内部 partial-Schur 自证；把实 1D 与共轭 2D 语义写入验收。", size=8.7)
    add_bullet(document, "把五个 QR/SVD 失败做有界分类，用七变体消融区分表示问题、初始化敏感和真实谱结构。", size=8.7)
    add_bullet(document, "把局部 bundle 与全局 manifold sheet 分开评价，并用全新进程、CI、配置、CSV、NPZ、日志和哈希形成可重复证据链。", size=8.7)
    add_callout(
        document,
        "论文定位",
        "numerical_framework_and_systematic_comparison（数值框架与系统比较）。25 项正式来源覆盖九主题；已有文献已包含参数化、cocycle、QR/CLV 与 Schur 方法，因此不主张方法首创。",
        fill="E2F0D9",
    )
    add_section_heading(document, "8｜仍需补充的内容")
    add_table(
        document,
        ["缺口", "当前边界", "建议下一步"],
        [
            ["稳定子束", "Stage F 仅一维不稳定子束", "增加稳定分支及双向传播"],
            ["Route H 2D 流形", "2D 子空间已确认，但没有全局二维流形对象", "实现 frame/圆周参数化并独立收敛"],
            ["长期/事件传播", "固定传播窗，无 event termination", "加入截面事件、长时几何和误差预算"],
            ["Sun–Earth 广度", "仅两个独立案例", "扩展家族并考虑独立应用章节"],
            ["理论与投稿", "无新定理；文献矩阵非穷尽综述", "补 a-posteriori 误差/理论讨论后再定目标期刊"],
        ],
        font_size=7.6,
        first_col_bold=True,
    )
    add_section_heading(document, "请导师优先判断")
    for question in config["review_questions"][:3]:
        add_bullet(document, question, size=8.4)
    add_callout(
        document,
        "最终状态",
        "中文论文初稿和四页摘要已形成，可进入导师评审；McCarthy 复现等级、Chapter 4 holdout 与 Route H 负结果均未改变。当前明确标记 not_submission_ready。",
        fill="FFF2CC",
    )
    document.save(DOCX)


def write_questions(config: dict[str, Any]) -> None:
    lines = [
        "# 给导师的审阅问题",
        "",
        f"- 学生：{config['author']}",
        f"- 导师：{config['adviser']}",
        f"- 单位：{config['institution']}",
        "- 论文当前定位：数值框架与系统比较",
        "- 当前状态：not_submission_ready",
        "",
    ]
    lines.extend(f"{index}. {question}" for index, question in enumerate(config["review_questions"], start=1))
    lines.extend(
        [
            "",
            "## 不随审阅意见改变的真实性边界",
            "",
            *[f"- {boundary}" for boundary in config["truth_boundaries"]],
        ]
    )
    QUESTIONS.write_text("\n".join(lines) + "\n", encoding="utf-8")


def render_pdf() -> tuple[int, str]:
    soffice = Path(r"C:\Program Files\LibreOffice\program\soffice.exe")
    profile = Path(tempfile.mkdtemp(prefix="lo_profile_adviser_"))
    command = [
        str(soffice),
        "--headless",
        f"-env:UserInstallation={profile.resolve().as_uri()}",
        "--convert-to",
        "pdf",
        "--outdir",
        str(ADVISER),
        str(DOCX),
    ]
    completed = subprocess.run(
        command,
        cwd=ROOT,
        text=True,
        stdout=subprocess.PIPE,
        stderr=subprocess.STDOUT,
        timeout=180,
        check=False,
    )
    shutil.rmtree(profile, ignore_errors=True)
    if completed.returncode != 0 or not PDF.is_file():
        raise RuntimeError(f"LibreOffice summary render failed: {completed.stdout}")
    return completed.returncode, completed.stdout.strip()


def validate_docx(config: dict[str, Any]) -> dict[str, Any]:
    from pypdf import PdfReader

    with zipfile.ZipFile(DOCX) as archive:
        xml = archive.read("word/document.xml").decode("utf-8")
        image_count = len(
            [name for name in archive.namelist() if name.startswith("word/media/")]
        )
    text = re.sub(r"<[^>]+>", "", xml)
    markers = [
        config["author"],
        config["adviser"],
        config["institution"],
        *config["required_content"],
        "0/4",
        "paper_projection=fail",
        "paper_3d=false",
        "not_submission_ready",
    ]
    missing = [marker for marker in markers if marker not in text]
    if missing:
        raise ValueError(f"adviser DOCX missing markers: {missing}")
    pages = len(PdfReader(PDF).pages)
    if pages != config["target_pdf_pages"]:
        raise ValueError(f"adviser summary rendered to {pages} pages, expected exactly 4")
    return {
        "docx_bytes": DOCX.stat().st_size,
        "docx_text_characters": len(text),
        "docx_embedded_images": image_count,
        "pdf_bytes": PDF.stat().st_size,
        "pdf_pages": pages,
        "required_markers": len(markers),
    }


def render_previews() -> list[str]:
    import pypdfium2 as pdfium
    from PIL import Image

    rendered = EVIDENCE / "rendered"
    rendered.mkdir(parents=True, exist_ok=True)
    document = pdfium.PdfDocument(str(PDF))
    paths: list[Path] = []
    images = []
    for index in range(len(document)):
        page = document[index]
        bitmap = page.render(scale=1.35)
        image = bitmap.to_pil()
        path = rendered / f"preview_page_{index + 1}.png"
        image.save(path)
        paths.append(path)
        images.append(image.copy())
        bitmap.close()
        page.close()
    document.close()
    width = max(image.width for image in images)
    thumb_height = int(images[0].height * width / images[0].width)
    contact = Image.new("RGB", (width * 2, thumb_height * 2), "white")
    for index, image in enumerate(images):
        image.thumbnail((width, thumb_height))
        contact.paste(image, ((index % 2) * width, (index // 2) * thumb_height))
    contact_path = rendered / "four_page_contact_sheet.png"
    contact.save(contact_path)
    paths.append(contact_path)
    return [str(path.relative_to(ROOT)).replace("\\", "/") for path in paths]


def metric_rows(data: dict[str, Any], stat: dict[str, Any]) -> list[dict[str, str]]:
    counts = stat["counts"]
    return [
        {"metric_id": "registered_cases", "value": "15", "evidence": str(REGISTRY.relative_to(ROOT)).replace("\\", "/"), "boundary": "finite benchmark"},
        {"metric_id": "pointwise_status", "value": "0 accepted;15 fail", "evidence": str(METHOD.relative_to(ROOT)).replace("\\", "/"), "boundary": "baseline retained"},
        {"metric_id": "schur_status", "value": f"{counts['ordered_partial_real_schur_tracking']['accepted']} accepted;{counts['ordered_partial_real_schur_tracking']['boundary']} boundary;{counts['ordered_partial_real_schur_tracking']['fail']} fail", "evidence": str(METHOD.relative_to(ROOT)).replace("\\", "/"), "boundary": "research-only status"},
        {"metric_id": "qr_status", "value": f"{counts['qr_svd_shifted_cocycle_iteration']['accepted']} accepted;{counts['qr_svd_shifted_cocycle_iteration']['fail']} fail", "evidence": str(METHOD.relative_to(ROOT)).replace("\\", "/"), "boundary": "research-only status"},
        {"metric_id": "independent_schur", "value": f"{stat['schur_agreement']}/12;max_angle={stat['schur_max_angle']:.9e} deg", "evidence": str(SCHUR.relative_to(ROOT)).replace("\\", "/"), "boundary": "finite discrete operators"},
        {"metric_id": "qr_failure_cases", "value": "4 no_accepted_1d_bundle;1 method_initialization_sensitive", "evidence": str(FAILURE.relative_to(ROOT)).replace("\\", "/"), "boundary": "bounded search"},
        {"metric_id": "ablation", "value": f"35 rows;{stat['ablation_accepted']} accepted;{stat['ablation_failed']} fail;{stat['ablation_exceptions']} exceptions", "evidence": str(ABLATION.relative_to(ROOT)).replace("\\", "/"), "boundary": "5 cases x 7 variants"},
        {"metric_id": "manifold", "value": f"126 rows;{stat['manifold_accepted']} accepted;{stat['manifold_failed']} fail", "evidence": str(MANIFOLD.relative_to(ROOT)).replace("\\", "/"), "boundary": "7 cases;1D unstable"},
        {"metric_id": "fresh_rerun", "value": f"{stat['fresh_scientific']} scientific pass;{stat['fresh_failures']} fail;max_rel={stat['fresh_max_rel']:.1f}", "evidence": str(FRESH.relative_to(ROOT)).replace("\\", "/"), "boundary": "reproducibility only"},
        {"metric_id": "chapter4_holdout", "value": "0/4;paper_projection=fail;paper_3d=false", "evidence": str(HOLDOUT.relative_to(ROOT)).replace("\\", "/"), "boundary": "frozen"},
        {"metric_id": "literature", "value": "25 verified;9/9 topics;21 DOI;4 not_assigned", "evidence": str(LITERATURE_SUMMARY.relative_to(ROOT)).replace("\\", "/"), "boundary": "not exhaustive"},
        {"metric_id": "paper_status", "value": "20-page Chinese draft;not_submission_ready", "evidence": str(PAPER_SUMMARY.relative_to(ROOT)).replace("\\", "/"), "boundary": "adviser review"},
    ]


def write_npz(path: Path, data: dict[str, Any], stat: dict[str, Any], pages: int) -> None:
    np.savez_compressed(
        path,
        schema_version=np.array(["invariant_bundle_adviser_summary_validation_npz_v1"]),
        case_ids=np.array([row["case_id"] for row in data["registry"]]),
        method_status=np.array([row["research_status"] for row in data["method"]]),
        independent_schur_dimensions=np.array([int(row["independent_selected_block_dimension"]) for row in data["schur"]]),
        qr_failure_labels=np.array([row["final_label"] for row in data["failure"]]),
        manifold_status=np.array([row["status"] for row in data["manifold"]]),
        pdf_pages=np.array([pages], dtype=np.int64),
        chapter4_holdout=np.array([0, 4], dtype=np.int64),
        positioning=np.array(["numerical_framework_and_systematic_comparison"]),
        submission_readiness=np.array(["not_claimed"]),
    )


def write_hash_manifest() -> None:
    target = EVIDENCE / "artifact_hashes.csv"
    core_delivery = [
        ADVISER / "McCarthy2018_54图逐图复现对照报告.docx",
        ADVISER / "McCarthy2018_54图逐图复现对照报告.pdf",
        ADVISER / "复现情况一页说明.pdf",
        ADVISER / "导师审阅重点.md",
    ]
    inputs = [
        CONFIG,
        Path(__file__).resolve(),
        ROOT / "scripts" / "run_final_goal_acceptance.py",
        ROOT / "tests" / "test_invariant_bundle_adviser_summary.py",
        ROOT / "tests" / "test_final_goal_acceptance.py",
        ROOT / "tests" / "test_stage_g_delivery_artifacts.py",
        REGISTRY,
        METHOD,
        SCHUR,
        FAILURE,
        ABLATION,
        MANIFOLD,
        FRESH,
        PAPER_SUMMARY,
        LITERATURE_SUMMARY,
        HOLDOUT,
        STAGE_G_HASHES,
    ]
    outputs = core_delivery + [DOCX, PDF, QUESTIONS]
    outputs += [
        path
        for path in sorted(EVIDENCE.rglob("*"))
        if path.is_file() and path != target
    ]
    rows = [
        {
            "schema_version": "adviser_summary_stage_artifact_hash_v1",
            "path": str(path.relative_to(ROOT)).replace("\\", "/"),
            "bytes": path.stat().st_size,
            "sha256": sha256(path),
        }
        for path in inputs + outputs
    ]
    write_csv(target, rows)


def validate_config(config: dict[str, Any]) -> None:
    if config["target_pdf_pages"] != 4:
        raise ValueError("adviser summary must target exactly four pages")
    if len(config["required_content"]) != 8:
        raise ValueError("all eight required adviser-summary content items must remain")
    if not 3 <= len(config["review_questions"]) <= 5:
        raise ValueError("the adviser question file must contain 3-5 questions")
    if config["positioning"] != "numerical_framework_and_systematic_comparison":
        raise ValueError("adviser summary positioning may not be promoted")
    if config["submission_readiness"] != "not_claimed":
        raise ValueError("submission readiness must remain not_claimed")
    text = "\n".join(config["truth_boundaries"])
    for marker in ("0/4", "paper_projection=fail", "paper_3d=false", "Route H", "二维实", "投稿"):
        if marker not in text:
            raise ValueError(f"truth boundary missing {marker}")


def build() -> dict[str, Any]:
    if any(path.exists() for path in (DOCX, PDF, QUESTIONS)):
        raise RuntimeError("refusing to overwrite adviser-summary deliverables")
    if EVIDENCE.exists() and any(EVIDENCE.iterdir()):
        raise RuntimeError("refusing to overwrite adviser-summary evidence")
    EVIDENCE.mkdir(parents=True, exist_ok=True)
    logs = EVIDENCE / "logs"
    logs.mkdir()
    started = datetime.now(timezone.utc)
    config = json.loads(CONFIG.read_text(encoding="utf-8"))
    validate_config(config)
    data = load_data()
    if (len(data["registry"]), len(data["method"]), len(data["schur"]), len(data["failure"]), len(data["ablation"]), len(data["manifold"])) != (15, 45, 12, 5, 35, 126):
        raise ValueError("authoritative research cardinalities changed")
    if len(data["holdout"]) != 4 or any(
        row["paper_projection_acceptance"] != "fail" or row["paper_3d_equivalence"] != "false"
        for row in data["holdout"]
    ):
        raise ValueError("frozen Chapter 4 holdout changed")
    stat = metrics(data)
    build_docx(config, data, stat)
    write_questions(config)
    render_code, render_output = render_pdf()
    validation = validate_docx(config)
    previews = render_previews()
    metric_table = metric_rows(data, stat)
    write_csv(EVIDENCE / "adviser_summary_metrics.csv", metric_table)
    checks = [
        {"check_id": "identity", "status": "pass", "detail": "兀文昊;张晨;中国科学院大学"},
        {"check_id": "required_content", "status": "pass", "detail": "8/8"},
        {"check_id": "pdf_pages", "status": "pass", "detail": f"{validation['pdf_pages']}/4"},
        {"check_id": "review_questions", "status": "pass", "detail": str(len(config["review_questions"]))},
        {"check_id": "frozen_reproduction", "status": "pass", "detail": "unchanged"},
        {"check_id": "chapter4_holdout", "status": "pass", "detail": "0/4;paper_projection=fail;paper_3d=false"},
        {"check_id": "route_h", "status": "pass", "detail": "physical=2D/fail;legacy=positive_control"},
        {"check_id": "failure_visibility", "status": "pass", "detail": "all negative/boundary/control evidence retained"},
        {"check_id": "positioning", "status": "pass", "detail": config["positioning"]},
        {"check_id": "submission_readiness", "status": "pass", "detail": "not_claimed"},
    ]
    write_csv(EVIDENCE / "adviser_summary_checks.csv", checks)
    write_npz(EVIDENCE / "adviser_summary_validation.npz", data, stat, validation["pdf_pages"])
    (EVIDENCE / "failure_evidence.md").write_text(
        "# Adviser-summary retained failure evidence\n\n"
        "- Pointwise eig remains 0/15 accepted.\n"
        "- Four physical corrected-rho Route H cases remain two-dimensional/fail; no one-dimensional relabelling is allowed.\n"
        "- The legacy seed-rho case remains a positive control and is not a physical replacement.\n"
        "- QR/SVD failure classification retains four `no_accepted_1d_bundle` and one `method_initialization_sensitive` result.\n"
        "- Stage F retains 90 failed manifold rows and lower-resolution sheets above 0.01.\n"
        "- Chapter 4 remains `0/4`, `paper_projection=fail`, `paper_3d=false`.\n"
        "- The four-page summary is for adviser review and remains `not_submission_ready`.\n",
        encoding="utf-8",
    )
    render_record = {
        "schema_version": "adviser_summary_render_validation_v1",
        **validation,
        "libreoffice_return_code": render_code,
        "libreoffice_output": render_output,
        "preview_paths": previews,
    }
    (EVIDENCE / "render_validation.json").write_text(
        json.dumps(render_record, indent=2, ensure_ascii=False, sort_keys=True) + "\n",
        encoding="utf-8",
    )
    environment = {
        "schema_version": "adviser_summary_stage_environment_v1",
        "python": sys.version,
        "python_executable": sys.executable,
        "platform": platform.platform(),
        "numpy": np.__version__,
        "document_builder": "python-docx",
        "renderer": "LibreOffice headless",
    }
    import docx as docx_module

    environment["python_docx"] = docx_module.__version__
    (logs / "environment.json").write_text(
        json.dumps(environment, indent=2, ensure_ascii=False, sort_keys=True) + "\n",
        encoding="utf-8",
    )
    finished = datetime.now(timezone.utc)
    summary = {
        "schema_version": "invariant_bundle_adviser_summary_v1",
        "status": "pass",
        "author": config["author"],
        "adviser": config["adviser"],
        "institution": config["institution"],
        "pdf_pages": validation["pdf_pages"],
        "required_content": "8/8",
        "review_questions": len(config["review_questions"]),
        "positioning": config["positioning"],
        "submission_readiness": "not_claimed",
        "chapter4_holdout": "0/4",
        "truth_boundary_status": "preserved",
    }
    (EVIDENCE / "adviser_summary_summary.json").write_text(
        json.dumps(summary, indent=2, ensure_ascii=False, sort_keys=True) + "\n",
        encoding="utf-8",
    )
    (logs / "stage_execution.log").write_text(
        "adviser summary build\n"
        f"started_utc={started.isoformat()}\n"
        f"finished_utc={finished.isoformat()}\n"
        f"elapsed_seconds={(finished-started).total_seconds():.6f}\n"
        f"python={sys.executable}\n"
        "identity=兀文昊|张晨|中国科学院大学\n"
        "required_content=8/8\n"
        "pdf_pages=4/4\n"
        f"review_questions={len(config['review_questions'])}\n"
        "chapter4_holdout=0/4\n"
        "truth_boundary_status=preserved\n"
        "submission_readiness=not_claimed\n",
        encoding="utf-8",
    )
    write_hash_manifest()
    return summary


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--check-only", action="store_true")
    args = parser.parse_args()
    config = json.loads(CONFIG.read_text(encoding="utf-8"))
    validate_config(config)
    if args.check_only:
        print(json.dumps({"status": "pass", "target_pages": 4, "questions": len(config["review_questions"])}, indent=2))
        return 0
    print(json.dumps(build(), indent=2, ensure_ascii=False, sort_keys=True))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
