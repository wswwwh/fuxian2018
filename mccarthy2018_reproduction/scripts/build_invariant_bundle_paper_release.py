#!/usr/bin/env python3
"""Build the complete Chinese invariant-bundle paper release and audit package."""

from __future__ import annotations

import argparse
import csv
from datetime import datetime, timezone
import hashlib
import json
from pathlib import Path
import platform
import re
import shutil
import statistics
import subprocess
import sys
import tempfile
from typing import Any, Iterable
import zipfile

import numpy as np


ROOT = Path(__file__).resolve().parents[1]
DEFAULT_CONFIG = (
    ROOT / "research" / "invariant_bundles" / "configs" / "paper_release.json"
)
DEFAULT_RELEASE = ROOT / "research" / "invariant_bundles" / "paper_release"
DEFAULT_EVIDENCE = (
    ROOT / "research" / "invariant_bundles" / "paper_release_validation"
)
PAPER = ROOT / "research" / "invariant_bundles" / "paper"
RESULTS = ROOT / "research" / "invariant_bundles" / "results"
REGISTRY = (
    ROOT / "research" / "invariant_bundles" / "benchmarks" / "benchmark_registry.csv"
)
METHOD_CSV = RESULTS / "csv" / "method_comparison.csv"
MANIFOLD_CSV = RESULTS / "csv" / "manifold_convergence.csv"
RESOLUTION_CSV = RESULTS / "csv" / "resolution_convergence.csv"
RUNTIME_CSV = RESULTS / "csv" / "runtime_scaling.csv"
SCHUR_CSV = RESULTS / "csv" / "independent_schur_backend_comparison.csv"
QR_FAILURE_CSV = RESULTS / "csv" / "qr_svd_failure_classification.csv"
ABLATION_CSV = RESULTS / "csv" / "ablation_study.csv"
LITERATURE_CONFIG = (
    ROOT
    / "research"
    / "invariant_bundles"
    / "configs"
    / "literature_verification.json"
)
LITERATURE_MATRIX = PAPER / "literature_matrix.csv"
VERIFIED_BIB = PAPER / "references_verified.bib"
FRESH_COMPARISON = (
    ROOT
    / "research"
    / "invariant_bundles"
    / "independent_rerun"
    / "comparison_to_stage_f.csv"
)
FRESH_REPORT = (
    ROOT
    / "research"
    / "invariant_bundles"
    / "independent_rerun"
    / "independent_rerun_report.md"
)
REPRO_SUMMARY = ROOT / "data" / "computed" / "reproduction_baseline_v1_summary.csv"
CHAPTER4_HOLDOUT = (
    ROOT
    / "data"
    / "computed"
    / "chapter4_fig43_fig46_projection_holdout_audit.csv"
)

METHODS = {
    "traditional_pointwise_eigendecomposition": "Pointwise eig",
    "ordered_partial_real_schur_tracking": "Partial real-Schur",
    "qr_svd_shifted_cocycle_iteration": "Shifted QR/SVD",
}
STATUS_ZH = {"accepted": "通过", "boundary": "边界", "fail": "失败"}
FAMILY_ZH = {
    "earth_moon_l1_quasi_halo": "地月 L1 拟 Halo",
    "earth_moon_l1_quasi_vertical": "地月 L1 拟 Vertical",
    "earth_moon_route_h_quasi_dro": "地月 Route H 拟 DRO",
    "sun_earth_l1_two_frequency_torus": "日地 L1 双频环面",
}
CLAIM_FIELDS = [
    "claim_id",
    "claim_text",
    "supporting_cases",
    "supporting_csv",
    "supporting_figure",
    "acceptance_threshold",
    "status",
    "limitation",
]


def read_csv(path: Path) -> list[dict[str, str]]:
    with path.open("r", encoding="utf-8", newline="") as stream:
        return list(csv.DictReader(stream))


def write_csv(path: Path, rows: list[dict[str, Any]], fields: list[str] | None = None) -> None:
    if not rows:
        raise ValueError(f"refusing to write empty CSV: {path}")
    path.parent.mkdir(parents=True, exist_ok=True)
    fieldnames = fields or list(rows[0])
    with path.open("w", encoding="utf-8", newline="") as stream:
        writer = csv.DictWriter(stream, fieldnames=fieldnames, extrasaction="ignore")
        writer.writeheader()
        writer.writerows(rows)


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as stream:
        for block in iter(lambda: stream.read(1024 * 1024), b""):
            digest.update(block)
    return digest.hexdigest().upper()


def sci(value: str | float, digits: int = 3) -> str:
    return f"{float(value):.{digits}e}"


def validate_config(config: dict[str, Any]) -> None:
    if config.get("study_positioning") != "numerical_framework_and_systematic_comparison":
        raise ValueError("paper positioning may not be promoted beyond systematic comparison")
    if config.get("submission_readiness") != "not_claimed":
        raise ValueError("submission readiness must remain not_claimed")
    if len(config.get("required_sections", [])) != 13:
        raise ValueError("the Chinese manuscript must retain all 13 required sections")
    boundary_text = "\n".join(config.get("truth_boundaries", []))
    for marker in (
        "0/4",
        "paper_projection=fail",
        "paper_3d=false",
        "Route H",
        "二维实",
        "投稿",
    ):
        if marker not in boundary_text:
            raise ValueError(f"truth boundary marker missing: {marker}")
    for figure in config.get("figures", []):
        source = ROOT / figure["source"]
        if not source.is_file():
            raise FileNotFoundError(source)


def load_data() -> dict[str, Any]:
    return {
        "registry": read_csv(REGISTRY),
        "method": read_csv(METHOD_CSV),
        "manifold": read_csv(MANIFOLD_CSV),
        "resolution": read_csv(RESOLUTION_CSV),
        "runtime": read_csv(RUNTIME_CSV),
        "schur": read_csv(SCHUR_CSV),
        "qr_failure": read_csv(QR_FAILURE_CSV),
        "ablation": read_csv(ABLATION_CSV),
        "literature": read_csv(LITERATURE_MATRIX),
        "fresh": read_csv(FRESH_COMPARISON),
        "repro": read_csv(REPRO_SUMMARY),
        "holdout": read_csv(CHAPTER4_HOLDOUT),
    }


def compute_metrics(data: dict[str, Any]) -> dict[str, Any]:
    method = data["method"]
    status_counts: dict[str, dict[str, int]] = {}
    for method_id in METHODS:
        status_counts[method_id] = {
            status: sum(
                row["method"] == method_id and row["research_status"] == status
                for row in method
            )
            for status in ("accepted", "boundary", "fail")
        }
    schur_angles = [
        float(row["invariant_subspace_principal_angle_max_deg"])
        for row in data["schur"]
    ]
    scientific = [
        row for row in data["fresh"] if row["comparison_kind"] != "informational_provenance"
    ]
    numeric = [
        row for row in data["fresh"] if row["comparison_kind"] == "numeric_scientific"
    ]
    numeric_rel = [float(row["relative_difference"]) for row in numeric]
    repro = {row["metric_id"]: row["value"] for row in data["repro"]}
    return {
        "status_counts": status_counts,
        "schur_max_angle_deg": max(schur_angles),
        "schur_agreement_rows": sum(
            row["dimension_agreement"] == "true"
            and row["classification_agreement"] == "true"
            and row["status_agreement"] == "true"
            for row in data["schur"]
        ),
        "manifold_accepted": sum(row["status"] == "accepted" for row in data["manifold"]),
        "manifold_failed": sum(row["status"] == "fail" for row in data["manifold"]),
        "max_jacobi_drift": max(float(row["manifold_jacobi_drift"]) for row in data["manifold"]),
        "max_initial_growth_deviation": max(
            abs(float(row["initial_linear_growth_ratio"]) - 1.0)
            for row in data["manifold"]
        ),
        "fresh_scientific_checks": len(scientific),
        "fresh_information_rows": len(data["fresh"]) - len(scientific),
        "fresh_failures": sum(row["comparison_status"] == "fail" for row in data["fresh"]),
        "fresh_max_relative_difference": max(numeric_rel) if numeric_rel else 0.0,
        "repro": repro,
        "reference_count": len(data["literature"]),
        "doi_count": sum(bool(row["doi"]) for row in data["literature"]),
    }


def status_cell(row: dict[str, str]) -> str:
    status = STATUS_ZH[row["research_status"]]
    dimension = row["bundle_dimension"] or "-"
    residual = row["max_invariance_residual"]
    return f"{status}; k={dimension}; r={sci(residual)}"


def build_summary_tables(
    release: Path,
    data: dict[str, Any],
    metrics: dict[str, Any],
) -> dict[str, list[dict[str, Any]]]:
    tables_dir = release / "tables"
    tables_dir.mkdir(parents=True, exist_ok=True)
    method_by_key = {(row["case_id"], row["method"]): row for row in data["method"]}
    case_rows: list[dict[str, Any]] = []
    for registry_row in data["registry"]:
        case_id = registry_row["case_id"]
        point = method_by_key[(case_id, "traditional_pointwise_eigendecomposition")]
        schur = method_by_key[(case_id, "ordered_partial_real_schur_tracking")]
        qr = method_by_key[(case_id, "qr_svd_shifted_cocycle_iteration")]
        case_rows.append(
            {
                "case_id": case_id,
                "family": FAMILY_ZH[registry_row["family"]],
                "spectral_samples": registry_row["spectral_samples"],
                "pointwise_status": point["research_status"],
                "pointwise_dimension": point["bundle_dimension"],
                "pointwise_max_residual": point["max_invariance_residual"],
                "schur_status": schur["research_status"],
                "schur_dimension": schur["bundle_dimension"],
                "schur_max_residual": schur["max_invariance_residual"],
                "qr_svd_status": qr["research_status"],
                "qr_svd_dimension": qr["bundle_dimension"],
                "qr_svd_max_residual": qr["max_invariance_residual"],
            }
        )
    write_csv(tables_dir / "table_15_case_summary.csv", case_rows)

    high_cases = [
        "em_halo_12p40_n45",
        "em_vertical_12p66_n57",
        "se_active_geometry_member_468",
    ]
    anchor_rows: list[dict[str, Any]] = []
    for case_id in high_cases:
        for method_id in METHODS:
            row = method_by_key[(case_id, method_id)]
            anchor_rows.append(
                {
                    "case_id": case_id,
                    "method": METHODS[method_id],
                    "bundle_dimension": row["bundle_dimension"],
                    "max_invariance_residual": row["max_invariance_residual"],
                    "phase_principal_angle_max_deg": row["phase_principal_angle_max_deg"],
                    "runtime_seconds": row["runtime_seconds"],
                    "research_status": row["research_status"],
                }
            )
    write_csv(tables_dir / "table_high_resolution_anchors.csv", anchor_rows)

    schur_rows = [
        {
            "case_id": row["case_id"],
            "internal_dimension": row["internal_selected_block_dimension"],
            "independent_dimension": row["independent_selected_block_dimension"],
            "principal_angle_max_deg": row["invariant_subspace_principal_angle_max_deg"],
            "classification_agreement": row["classification_agreement"],
            "status_agreement": row["status_agreement"],
            "validation_verdict": row["validation_verdict"],
        }
        for row in data["schur"]
    ]
    write_csv(tables_dir / "table_independent_schur_validation.csv", schur_rows)

    failure_rows = [
        {
            "case_id": row["case_id"],
            "independent_schur_dimension": row["independent_schur_dimension"],
            "best_native_status": row["best_native_research_status"],
            "best_native_max_residual": row["best_native_max_invariance_residual"],
            "iteration_caps": row["iteration_caps_tested"],
            "resolutions": row["resolutions_tested"],
            "high_precision_digits": row["high_precision_decimal_digits"],
            "final_label": row["final_label"],
            "negative_result_retained": row["negative_result_retained"],
        }
        for row in data["qr_failure"]
    ]
    write_csv(tables_dir / "table_qr_svd_failure_classification.csv", failure_rows)

    ablation_rows: list[dict[str, Any]] = []
    for variant in dict.fromkeys(row["variant"] for row in data["ablation"]):
        subset = [row for row in data["ablation"] if row["variant"] == variant]
        finite = [
            float(row["bundle_residual_max"])
            for row in subset
            if row["bundle_residual_max"].lower() != "nan"
        ]
        ablation_rows.append(
            {
                "variant": variant,
                "short_label": subset[0]["variant_short_label"],
                "accepted": sum(row["research_status"] == "accepted" for row in subset),
                "failed": sum(row["research_status"] == "fail" for row in subset),
                "method_exceptions": sum(row["bundle_dimension"] == "0" for row in subset),
                "median_finite_max_residual": statistics.median(finite),
                "max_finite_max_residual": max(finite),
            }
        )
    write_csv(tables_dir / "table_ablation_variant_summary.csv", ablation_rows)

    manifold_rows: list[dict[str, Any]] = []
    for case_id in dict.fromkeys(row["case_id"] for row in data["manifold"]):
        for method_id in METHODS:
            subset = [
                row
                for row in data["manifold"]
                if row["case_id"] == case_id and row["method"] == method_id
            ]
            if not subset:
                continue
            cross_values = [
                float(row["cross_resolution_normalized_3d_distance"])
                for row in subset
                if row["cross_resolution_normalized_3d_distance"]
                and row["cross_resolution_normalized_3d_distance"].lower() != "nan"
            ]
            manifold_rows.append(
                {
                    "case_id": case_id,
                    "method": METHODS[method_id],
                    "rows": len(subset),
                    "accepted": sum(row["status"] == "accepted" for row in subset),
                    "failed": sum(row["status"] == "fail" for row in subset),
                    "max_jacobi_drift": max(float(row["manifold_jacobi_drift"]) for row in subset),
                    "max_direction_angle_to_qr_deg": max(float(row["direction_principal_angle_max_deg_to_qr"]) for row in subset),
                    "max_cross_resolution_distance": max(cross_values) if cross_values else "",
                }
            )
    write_csv(tables_dir / "table_manifold_case_summary.csv", manifold_rows)

    runtime_rows = [
        {
            "case_id": row["case_id"],
            "method": METHODS[row["method"]],
            "spectral_samples": row["spectral_samples"],
            "runtime_seconds": row["runtime_seconds"],
            "iterations": row["iterations"],
            "research_status": row["research_status"],
        }
        for row in data["runtime"]
        if row["case_id"] in high_cases
    ]
    write_csv(tables_dir / "table_runtime_anchor_summary.csv", runtime_rows)

    repro_keys = [
        "target_rows",
        "v0_targets",
        "v2_targets",
        "evidence_accepted",
        "evidence_boundary",
        "evidence_diagnostic",
        "evidence_proxy",
        "chapter4_frozen_holdout_pass",
        "chapter4_frozen_holdout_total",
    ]
    repro_rows = [
        {
            "metric_id": key,
            "value": metrics["repro"][key],
            "interpretation": next(
                row["interpretation"] for row in data["repro"] if row["metric_id"] == key
            ),
        }
        for key in repro_keys
    ]
    write_csv(tables_dir / "table_reproduction_truth_boundary.csv", repro_rows)
    return {
        "cases": case_rows,
        "anchors": anchor_rows,
        "schur": schur_rows,
        "failures": failure_rows,
        "ablation": ablation_rows,
        "manifold": manifold_rows,
        "runtime": runtime_rows,
        "repro": repro_rows,
    }


def copy_figures(config: dict[str, Any], release: Path) -> list[dict[str, str]]:
    target_dir = release / "figures"
    target_dir.mkdir(parents=True, exist_ok=True)
    copied: list[dict[str, str]] = []
    for index, item in enumerate(config["figures"], start=1):
        source = ROOT / item["source"]
        target = target_dir / source.name
        shutil.copy2(source, target)
        vector_source = source.with_suffix(".pdf")
        if vector_source.is_file():
            shutil.copy2(vector_source, target_dir / vector_source.name)
        copied.append(
            {
                "figure_number": str(index),
                "file": target.name,
                "caption": item["caption"],
                "source_sha256": sha256(source),
            }
        )
    write_csv(target_dir / "figure_manifest.csv", copied)
    return copied


def md_escape(value: Any) -> str:
    return str(value).replace("|", "\\|").replace("\n", " ")


def md_table(rows: Iterable[dict[str, Any]], columns: list[tuple[str, str]]) -> str:
    rows = list(rows)
    header = "| " + " | ".join(label for _, label in columns) + " |"
    separator = "| " + " | ".join("---" for _ in columns) + " |"
    body = [
        "| " + " | ".join(md_escape(row[key]) for key, _ in columns) + " |"
        for row in rows
    ]
    return "\n".join([header, separator, *body])


def compact_case_rows(data: dict[str, Any]) -> list[dict[str, str]]:
    method_by_key = {(row["case_id"], row["method"]): row for row in data["method"]}
    result: list[dict[str, str]] = []
    for registry_row in data["registry"]:
        case_id = registry_row["case_id"]
        result.append(
            {
                "case": case_id,
                "N": registry_row["spectral_samples"],
                "Pointwise": status_cell(method_by_key[(case_id, "traditional_pointwise_eigendecomposition")]),
                "Schur": status_cell(method_by_key[(case_id, "ordered_partial_real_schur_tracking")]),
                "QR": status_cell(method_by_key[(case_id, "qr_svd_shifted_cocycle_iteration")]),
            }
        )
    return result


def reference_list(literature_config: dict[str, Any]) -> str:
    lines = []
    for index, row in enumerate(literature_config["references"], start=1):
        doi = f" DOI: {row['doi']}." if row["doi"] else " DOI：未分配。"
        lines.append(
            f"{index}. {row['authors']}. {row['title']}. {row['venue']}, {row['year']}.{doi} {row['official_url']}"
        )
    return "\n\n".join(lines)


def manuscript_markdown(
    config: dict[str, Any],
    data: dict[str, Any],
    metrics: dict[str, Any],
) -> str:
    counts = metrics["status_counts"]
    point = counts["traditional_pointwise_eigendecomposition"]
    schur = counts["ordered_partial_real_schur_tracking"]
    qr = counts["qr_svd_shifted_cocycle_iteration"]
    method_by_key = {(row["case_id"], row["method"]): row for row in data["method"]}
    anchors = [
        "em_halo_12p40_n45",
        "em_vertical_12p66_n57",
        "se_active_geometry_member_468",
    ]
    anchor_labels = {
        "em_halo_12p40_n45": "Halo N45",
        "em_vertical_12p66_n57": "Vertical N57",
        "se_active_geometry_member_468": "Sun–Earth 468",
    }
    anchor_display = []
    for case_id in anchors:
        anchor_display.append(
            {
                "case": anchor_labels[case_id],
                "point": sci(method_by_key[(case_id, "traditional_pointwise_eigendecomposition")]["max_invariance_residual"]),
                "schur": sci(method_by_key[(case_id, "ordered_partial_real_schur_tracking")]["max_invariance_residual"]),
                "qr": sci(method_by_key[(case_id, "qr_svd_shifted_cocycle_iteration")]["max_invariance_residual"]),
                "schur_time": f"{float(method_by_key[(case_id, 'ordered_partial_real_schur_tracking')]['runtime_seconds']):.3f}",
                "qr_time": f"{float(method_by_key[(case_id, 'qr_svd_shifted_cocycle_iteration')]['runtime_seconds']):.3f}",
            }
        )
    failure_display = [
        {
            "case": row["case_id"],
            "k": row["independent_schur_dimension"],
            "best": STATUS_ZH[row["best_native_research_status"]],
            "r": sci(row["best_native_max_invariance_residual"]),
            "label": row["final_label"],
        }
        for row in data["qr_failure"]
    ]
    literature_config = json.loads(LITERATURE_CONFIG.read_text(encoding="utf-8"))
    compact = compact_case_rows(data)
    return f"""# {config['title']}

- 作者：{config['author']}
- 导师：{config['adviser']}
- 单位：{config['institution']}
- 日期：{config['date']}
- 稿件定位：数值框架与系统比较（`numerical_framework_and_systematic_comparison`）
- 状态：导师评审初稿；`not_submission_ready`；不声明已经达到投稿条件

## 摘要

拟周期轨道附近的稳定与不稳定方向并非若干相位点上彼此独立的特征向量，而是由相位平移线性 cocycle 共同定义的实不变子束。若直接对每个局部状态转移矩阵做点式特征分解，容易出现方向符号跳变、复特征向量取实部后失去不变性、分辨率变化时方向不收敛，以及局部方向误差被放大为全局流形片几何误差。本文以冻结的 McCarthy 2018 复现工程为数据边界，构建 15 个案例、4 类轨道族、3 种方法的可审计 benchmark，对传统点式 eig、partial real-Schur 子空间选择和 shifted QR/SVD cocycle 迭代进行统一比较。

在固定阈值下，点式 eig 为 {point['accepted']}/15 通过，partial real-Schur 为 {schur['accepted']}/15 通过、{schur['boundary']}/15 边界、{schur['fail']}/15 失败，shifted QR/SVD 为 {qr['accepted']}/15 通过、{qr['fail']}/15 失败。MATLAB R2024a 的独立 `schur/ordschur` 后端在 12 个关键案例上实现维数、分类和状态 {metrics['schur_agreement_rows']}/12 一致，最大子空间主角为 {metrics['schur_max_angle_deg']:.3e}°，低于 1e-4° 门槛。对 QR/SVD 的五个失败案例进行三种初始化、200/500/1000 次迭代、原生 N45 与 Fourier N67 诊断以及 80 位残差重算后，四个 physical corrected-rho Route H 案例仍是二维实共轭子空间上的一维失败；legacy seed-rho 案例仅在 Schur 初始化下通过，被分类为方法初始化敏感。

流形阶段保存 126 行全状态传播证据，其中 {metrics['manifold_accepted']} 行通过、{metrics['manifold_failed']} 行失败；最大 Jacobi 漂移为 {metrics['max_jacobi_drift']:.3e}。高分辨率 Halo N45、Vertical N57 和 Sun–Earth 468 上，两种改进方法同时通过局部子束与测试范围内的流形门槛，但低分辨率全流形片到最高分辨率参考的距离仍超过 0.01。因此本文的可辩护贡献是一个带独立后端、失败分类、消融、全新进程重跑和 CI 守护的数值比较框架，而非全新理论。冻结的 Chapter 4 projection holdout 仍为 0/4，`paper_projection=fail`、`paper_3d=false`，本文不据此声称整篇学位论文等价复现或投稿就绪。

**关键词：** 拟周期轨道；不变子束；线性 cocycle；实 Schur 分解；QR/SVD；协变 Lyapunov 向量；CR3BP；不变流形；可重复性

## 1. 引言

拟周期运动是周期轨道与一般非周期运动之间的重要中间结构。在圆型限制性三体问题（CR3BP）中，周期轨道附近常存在双频或更高维拟周期环面，其内部相位演化可由旋转数表示，法向稳定性则决定扰动是收缩、扩张还是保持近中性。Olikara 与 Howell 的 Fourier/多重打靶计算以及 McCarthy 2018 的 stroboscopic mapping 工作构成了本研究的直接应用背景 [OlikaraHowell2010; Olikara2010Thesis; McCarthy2018]。后续工作进一步把拟周期环面及其流形用于四体模型与异宿连接 [McCarthyHowell2021; McCarthyHowell2023; OwenBaresi2024]。

数值困难集中在“方向”究竟是什么。对单个矩阵而言，特征向量满足局部代数方程；对拟周期环面而言，法向对象却必须同时满足相位平移和线性传播。Jorba 对不变曲线法向行为的计算、Haro 与 de la Llave 的参数化方法、Wysham 与 Meiss 的 cocycle 迭代以及 Huguet 等人的快速子束迭代均表明，目标是沿整条相位曲线耦合的对象 [Jorba2001; HaroLlave2006Numerical; WyshamMeiss2006; HuguetLlaveSire2013]。因此，“每个相位都找到一个很精确的局部特征向量”并不等价于“找到一个 cocycle 不变子束”。

本文从 54 图 McCarthy 复现工程中引出这一问题，但严格区分复现层和研究层。复现层已实现 54 个工程目标（13 个 V0、41 个 V2），证据等级仍为 7 个 accepted、30 个 boundary、5 个 diagnostic、12 个 proxy；这些标签被冻结。研究层只能读取其轨道、状态和元数据，不能把独立子束实验的成功回写为论文等价。尤其 Chapter 4 冻结 projection holdout 为 0/4，任何后验候选轨道或研究图都不得改变该结论。

依据 25 项核实的正式来源，partial real-Schur、连续正交化、QR/SVD 谱诊断、协变 Lyapunov 向量以及流形参数化都已有明确先例 [DieciRussellVanVleck1994; DieciVanVleck2002; GinelliEtAl2007; KuptsovParlitz2012; BaiDemmel1993; GranatKagstrom2006; HaroEtAl2016]。本文不使用“首次提出”或“全新理论”等措辞。贡献限定为：建立统一 benchmark；在相同源轨道和门槛下比较三类数值对象；用真正独立的 MATLAB Schur 后端复核；保留五个失败案例和七变体消融；把局部子束与全局流形片分开验收；通过全新进程重跑和 GitHub Actions 约束可重复性。

## 2. 问题描述与 cocycle 方程

设拟周期不变曲线由相位 θ∈T 参数化，离散映射使 θ 前进 ρ。沿曲线的线性化状态转移矩阵记为 A(θ)。秩为 k 的实子束由正交基 E(θ)∈R^(6×k) 表示，它必须满足

$$A(θ)E(θ) ≈ E(θ+ρ)R(θ),$$

其中 R(θ) 是子束内的约化映射；k=1 时为实标量，k=2 时为 2×2 实矩阵。本文采用归一化 Frobenius 残差

$$r(θ_i)=||A_iE_i-E_(i+ρ)(E_(i+ρ)^T A_iE_i)||_F / max(||A_iE_i||_F, ε_machine),$$

并同时报告最大/平均残差、相邻相位主角、跨分辨率主角、谱与倒数配对误差、复部相对量、运行时间、迭代次数和状态。通过阈值固定为 max r≤1e-6；1e-6<max r≤1e-3 仅可进入 boundary，若另有维数、源轨道或收敛失败则仍判 fail。

点式 eig 解的是 A(θ_i)v_i=λ_i v_i，而 cocycle 子束要求 A(θ_i)E_i 的像落入 θ_i+ρ 处的子空间。这两个方程只有在额外的可约化、分支一致和相位匹配条件成立时才可能一致。Eliasson 的几乎可约化理论也说明，可约化是需要条件的动力学性质，不能从每个相位都可对角化自动推出 [Eliasson2001]。这回答第一个科学问题：**pointwise eig 不是 cocycle invariant bundle，因为它没有约束相位平移后的像空间。**

第二个关键问题是实结构。若选中谱是一对 a±ib，任何单个复特征向量都不属于一维实空间；它的实部与虚部共同张成二维实不变子空间。Bai 与 Demmel 对实 Schur 1×1/2×2 块的重排以及 Granat 与 Kågström 对周期矩阵乘积 Schur 块的处理给出标准线性代数背景 [BaiDemmel1993; GranatKagstrom2006]。因此把复向量简单取实部会改变不变对象；**复共轭对不能被投影或重命名为一维实方向。**

## 3. 传统点式特征方向方法及其失效模式

基线方法在每个 θ_i 上独立分解 A_i，按模长选择稳定或不稳定分支；若得到复向量，则取实部、归一化，再做相邻符号对齐。这一实现不是为了制造弱基线，而是刻意保留工程中常见的捷径，以便区分局部特征对残差和全局 cocycle 残差。

15 个案例上，点式 eig 的研究状态全部为 fail。三个高分辨率锚点 Halo N45、Vertical N57、Sun–Earth 468 的最大 cocycle 残差分别为 {anchor_display[0]['point']}、{anchor_display[1]['point']} 和 {anchor_display[2]['point']}。消融进一步显示，仅增加符号对齐可把 Halo N45 残差从 1.616e-1 降到 1.240e-1，却仍高于 1e-6 通过阈值约五个数量级；Vertical N57 与 Sun–Earth 468 也保持约 1e-1 量级。

这类失败并非“特征分解不够精确”。局部矩阵的特征对可以达到很小代数残差，但不同相位选出的向量没有保证属于同一全局分支；谱接近、符号翻转、复共轭对和网格插值都会造成跨相位不一致。协变 Lyapunov 向量文献进一步区分了正交 QR 向量、奇异向量和真正协变方向 [GinelliEtAl2007; KuptsovParlitz2012]。因此本文不把任一单步 QR 列或局部 SVD 右奇异向量直接称为物理不变方向。

![三种方法总体结果](figures/fig_bundle_method_summary.png)

![代表性案例相位连续性](figures/fig_phase_continuity_profiles.png)

## 4. Partial real-Schur 方法

第一种改进方法先构造谱配点算子 G_N=(P_(θ+ρ→θ)⊗I_6)diag(A_0,…,A_(N-1))，再选择目标实 Schur 块。若目标根在冻结容差内为实数，取 k=1；若为共轭复根，则以实部和虚部构造 k=2 的实基，绝不把二维块改写成一维。节点基经局部 QR 正交化与相位子空间对齐后，再回到点式 cocycle 方程计算科学残差。这里“partial”表示只提取并验证所选不变块，不意味着本文提出新的 Schur 理论。

内部 Python 路线在锁定的 SciPy 运行时不能直接调用完整 `scipy.linalg.schur`，因此使用有序目标特征对构造并验证实 partial-Schur 块。为消除同源实现的自证风险，阶段 2 把 12 个关键 G_N 算子导出给 MATLAB R2024a `schur/ordschur`。独立后端采用 Intel oneAPI MKL 2023.2；12/12 案例的块维数、分类与研究状态一致，最大主角 {metrics['schur_max_angle_deg']:.3e}°，低于 1e-4° 门槛。该结果证明内部选中子空间与独立实 Schur 子空间在测试范围内一致，但不构成算法收敛定理。

在 15 案例中，partial real-Schur 得到 {schur['accepted']} 个 accepted、{schur['boundary']} 个 boundary、{schur['fail']} 个 fail。三个锚点的最大残差依次为 {anchor_display[0]['schur']}、{anchor_display[1]['schur']}、{anchor_display[2]['schur']}。Route H physical corrected-rho 的四个案例则均由独立后端确认为 k=2 共轭块，并保持 fail。这说明 Schur 方法解决的是“实子空间维数与谱块选择”问题，不保证任意选中块都满足低残差的一维物理 bundle。

## 5. Shifted QR/SVD cocycle 迭代

第二种改进路线从局部右奇异子空间初始化，反复执行 A_iE_i 的传播、在 θ_i+ρ 网格上的 QR、插值回基准网格、再次 QR，并对当前帧与上一迭代/相邻相位做符号或 Procrustes 子空间对齐。停止条件是最大子空间更新角不超过 2e-6°，默认上限 200 次。Dieci、Russell 与 Van Vleck 的连续正交化以及 Dieci 与 Van Vleck 的 QR/SVD 谱计算提供了方法学背景 [DieciRussellVanVleck1994; DieciVanVleck2002]。

QR/SVD 在 15 案例中 {qr['accepted']} 个 accepted、{qr['fail']} 个 fail，三个锚点残差分别为 {anchor_display[0]['qr']}、{anchor_display[1]['qr']} 和 {anchor_display[2]['qr']}。与 Schur 相比，它直接通过相位平移迭代逼近子空间，在若干一维案例上残差更低；但它对初始化、迭代上限和目标维数敏感，尤其不能自己把一个本质二维的共轭块变成一维实方向。

Schur 与 QR/SVD 解决的问题不同。Schur 路线负责稳定地识别实谱块、确定 k=1 还是 k=2，并为复杂谱提供可审计子空间；QR/SVD 路线负责在 cocycle 传播下迭代相位依赖帧，并提供收敛历史和 SVD 诊断。二者互为交叉检查而非简单替代。若 Schur 确认 k=2，而研究问题只接受一维分支，则正确结果是保留二维对象并报告一维失败，而不是强制降维。

## 6. Benchmark 轨道族与评价指标

注册表包含 15 个案例和四类轨道族：5 个地月 L1 拟 Halo、3 个地月 L1 拟 Vertical、4 个 Route H physical corrected-rho 加 1 个 legacy seed-rho 正控制，以及 2 个日地 L1 双频环面案例。谱网格从 N9 到 N57；Halo 与 Vertical 提供跨分辨率序列，Route H 提供算子语义与复谱压力测试，Sun–Earth 提供跨系统验证。所有案例均以状态文件哈希、映射时间、ρ、质量参数、积分器和源门槛登记，运行过程不得盲目延伸轨道族来追求更高通过率。

表 1 给出全部 15 案例的状态、实子空间维数 k 和最大不变性残差 r。accepted、boundary、fail 是研究层标签，不能写回冻结的 54 图复现等级。

{md_table(compact, [('case','案例'),('N','N'),('Pointwise','Pointwise eig'),('Schur','Partial Schur'),('QR','Shifted QR/SVD')])}

评价分为四层。第一层是源曲线/映射闭合；第二层是局部 bundle 的维数、残差、相位连续性和跨分辨率主角；第三层是传播后 Jacobi 漂移、线性增长一致性和相对 QR 的归一化位移距离；第四层是跨分辨率全流形片距离。只有前一层通过，后一层结果才有解释意义；但前一层通过并不自动使后一层通过。

![分辨率序列的子束主角收敛](figures/fig_resolution_convergence.png)

## 7. 子束结果

总体结果首先否定了“局部特征方向足够好”的假设：Pointwise eig 0/15 accepted。partial real-Schur 与 shifted QR/SVD 在 Halo、Vertical 和 Sun–Earth 锚点上把残差从约 1e-1 降到 1e-7 至 1e-12 区间。两种改进方法在同一锚点的方向差约 4e-5°或更小，说明当一维实子束存在且数值可分离时，两条路线会收敛到一致几何对象。

表 2 汇总三个高分辨率案例的残差与运行时间。

{md_table(anchor_display, [('case','案例'),('point','Pointwise r_max'),('schur','Schur r_max'),('qr','QR/SVD r_max'),('schur_time','Schur 秒'),('qr_time','QR/SVD 秒')])}

独立 MATLAB 后端的价值不在于再产生一张“看起来相同”的图，而在于验证实子空间分类。它对 Halo N21/N33/N45、Vertical N33/N45/N57、Sun–Earth 468、Route H 17/32/54/68 以及 legacy 68 共 12 案例给出 12/12 一致。特别地，四个 physical Route H 均为二维/失败，legacy 68 为一维/通过。这个结果阻止了用内部实现偏好解释 Route H 负结果。

全新进程重跑进一步从空 cocycle 缓存生成 15 个 cocycle、45 行 bundle 和 126 行 manifold 结果。与 Stage F 权威表逐字段比较 6156 项，其中 {metrics['fresh_scientific_checks']} 项科学检查全部通过、{metrics['fresh_information_rows']} 项运行 ID/时间/哈希等来源字段保留为信息，失败为 {metrics['fresh_failures']}，科学数值最大相对差为 {metrics['fresh_max_relative_difference']:.1f}。这验证实现稳定性，但不提升任何冻结复现标签。

## 8. 消融实验

消融覆盖 5 个案例×7 个变体，共 35 行：点式 eig 无相位对齐、仅符号对齐；Schur 无/有相位跟踪；QR/SVD 无/有相位对齐；以及 QR/SVD 使用 Schur 维数种子。总计 15 行 accepted、20 行 fail，并保留 4 个点式方法异常，异常不从分母中删除。

第一，点式符号对齐改善表面连续性但不能修复 cocycle 残差，说明符号跳变只是问题的一部分。第二，在三个一维锚点上，Schur 相位跟踪前后残差比约为 1，QR/SVD 相位对齐前后残差也约为 1；当子束已清晰分离时，相位对齐主要稳定表示而不是改变物理子空间。第三，在 Route H 2D 案例上，Schur 子空间相位跟踪可把 member 68 残差从 8.126e-1 降到 1.650e-1、member 32 从 6.537e-1 降到 4.279e-1，但仍是 fail。第四，给 QR/SVD 注入 Schur 维数能恢复“二维实对象”的语义，却没有把 Route H 变成可接受的一维 bundle。

![消融实验：子束残差](figures/ablation_bundle_residual.png)

![消融实验：相位连续性](figures/ablation_phase_continuity.png)

![消融实验：线性化一映射几何诊断](figures/ablation_manifold_geometry.png)

最后一幅图的 `manifold_geometry_distance` 仅来自线性化一映射点云，是消融诊断，不等价于 Stage F 的非线性全状态流形片，也不能替代后者的 0.01 跨分辨率门槛。

## 9. 流形传播与几何收敛

Stage F 选择 7 个案例、3 种方法、3 个全状态扰动范数（5e-8、1e-7、2e-7）、正负两个方向和 41 个时间采样，形成 126 行结果。不同方法共享源状态、相位、传播时长、DOP853 容差、坐标系和停止规则。最大 Jacobi 漂移为 {metrics['max_jacobi_drift']:.3e}，初始线性增长比相对 1 的最大偏差为 {metrics['max_initial_growth_deviation']:.3e}。

Halo N45、Vertical N57、Sun–Earth 468 上，Schur 与 QR/SVD 在所有测试扰动和两个符号下均通过，合计 {metrics['manifold_accepted']} 行 accepted。点式方法因上游 bundle 残差不合格而全部失败；Route H 物理案例也因目标一维 bundle 不成立而保持失败。三锚点上，点式方向相对 QR 相差约 7–9°，Schur 相对 QR 约 4e-5°以内；按扰动幅值归一化后，点式到 QR 的流形位移片距离约为 1e-2，而 Schur 到 QR 为 1e-7 或更小。

![流形传播指标](figures/fig_manifold_method_metrics.png)

![Halo 高分辨率归一化位移流形片](figures/fig_halo_manifold_displacement_sheets.png)

然而局部 bundle 收敛与全局 manifold sheet 收敛必须分开评价。Halo N21/N33 相对 N45 的全片距离约为 0.0219/0.0150，Vertical N33/N45 相对 N57 为 0.0245/0.0195，全部高于 0.01。局部主角可以已经很小，但有限时间传播、非线性曲率、初值离散和相位采样会积累成全片几何差异。因此本文只能说三个高分辨率锚点在测试传播窗内一致，不能说所有分辨率的全局流形已收敛。

## 10. Route H 算子语义案例

Route H 是本文最重要的负结果。legacy member 68 使用 seed-rho，选出近实一维正控制，但源曲线映射残差约 1.988e-3；physical corrected-rho 把映射闭合改进到约 8.697e-13，却显示目标谱为相对复部约 0.342 的共轭对。正确物理算子因此要求二维实子空间，而不是更“漂亮”的一维方向。映射闭合改善约九个数量级却失去一维通过，不是回退到旧 ρ 的理由，而是算子语义决定结果的证据。

![Route H corrected-rho 与 legacy seed-rho 对照](figures/fig_route_h_rho_control.png)

对五个 QR/SVD 失败案例的有界分类如下。

{md_table(failure_display, [('case','案例'),('k','独立 Schur k'),('best','最佳状态'),('r','最佳 r_max'),('label','最终分类')])}

member 17/32/54/68 的分支选择与独立 Schur 一致；三种初始化、三个迭代上限和 N67 Fourier lift 均未得到 accepted 一维 bundle，80 位 mpmath 只重算残差而没有伪装成任意精度轨迹积分。故其标签是 `no_accepted_1d_bundle`，而不是“算得不够久”。legacy 68 从 Schur seed 可通过、从 random/local-SVD 不稳定，被标为 `method_initialization_sensitive`。所有负结果原样进入 CSV、NPZ 和失败证据。

## 11. 计算成本

三个高分辨率锚点上，partial real-Schur 用时 {anchor_display[0]['schur_time']}、{anchor_display[1]['schur_time']}、{anchor_display[2]['schur_time']} 秒，QR/SVD 用时 {anchor_display[0]['qr_time']}、{anchor_display[1]['qr_time']}、{anchor_display[2]['qr_time']} 秒。二者在这些可接受案例上都远低于每案例 wall-time 上限。Pointwise eig 最快，但其 0/15 accepted 使单纯速度比较没有科学意义。

Route H 的失败 QR 运行会达到 200 次迭代上限并耗时数秒；失败分类又扩展到 500/1000 次、三种初始化和 N67 诊断。因此成本必须与状态共同报告。把失败行从运行时间统计中删除，会系统性低估鲁棒算法在困难谱结构上的真实代价。

实现使用 `D:\\miniconda3\\envs\\cislunar\\python.exe`、Python 3.11、NumPy/OpenBLAS 和 SciPy；独立 Schur 使用 MATLAB R2024a 24.1.0.2537033 与 Intel oneAPI MKL 2023.2。全新进程重跑记录控制器和工作进程 PID、命令、缓存语义、环境、CSV/NPZ 哈希。GitHub Actions 快速 CI 在 push/PR 上运行导入、单元测试、注册表、小型物理 benchmark、文档一致性和权威哈希守护；完整研究验证由手动 workflow_dispatch 触发并把输出写到 runner 临时目录。

## 12. 局限性与讨论

第一，本文不是整篇 McCarthy 2018 的严格数值等价复现。54 图工程覆盖和 25 项外部文献核实只是研究起点；Chapter 4 projection holdout 仍为 0/4，`paper_projection=fail`、`paper_3d=false`。后验 12.397983 日 N21 候选只能作为根因线索，不能替换冻结 holdout。

第二，本文没有给出子束存在、唯一性、可约化性或收敛率的新证明。Haro 与 de la Llave 的严格结果所需假设和 a-posteriori 估计超出本研究 [HaroLlave2006Rigorous]。MATLAB 独立后端验证了 12 个有限离散算子的实 Schur 子空间一致性，不等于连续问题定理，也不是第三个库/硬件后端的交叉验证。

第三，QR/SVD 失败分类是有界搜索。迭代上限为 1000，诊断分辨率至 N67；80 位计算只重算残差，没有用任意精度重积分 CR3BP 轨迹。因而 `no_accepted_1d_bundle` 的精确含义是“在声明的配置空间内未发现通过的一维 bundle”，而不是对所有算法和所有分辨率的不可能性证明。

第四，流形范围限于 7 案例的一维不稳定子束、固定传播窗、三种扰动、两个符号、无事件终止和 CR3BP 会合坐标。稳定子束、二维 Route H 流形对象、长时间事件传播、更多 Sun–Earth 案例和更高保真星历模型尚未独立验证。局部 bundle 残差与全局流形片几何是不同验收对象，不能互相代替。

第五，文献矩阵覆盖九个指定主题和 25 项正式来源，但不是系统综述或穷尽性 novelty search。21 个 DOI 已核实，四个学位/会议来源明确记录为未分配 DOI；页面访问受 robots 或付费墙限制时使用出版社元数据、DOI 注册和机构库交叉核验。因此当前定位只能是“数值框架与系统比较”，不应升级为方法首创。

最后，独立后端、失败分类、消融、全新进程和 CI 使证据链更可靠，却不自动满足期刊的理论深度、统计广度、稳定子束、二维流形或外部数据要求。本文是可交导师评审的完整初稿，**不声明已经达到投稿条件**。

## 13. 结论

本文回答了七个核心科学问题。其一，pointwise eig 只满足局部代数方程，不满足相位平移 cocycle 方程，故不是自动成立的不变子束。其二，复共轭对在实数域对应二维 Schur 子空间，不能投影为一维实方向。其三，Schur 解决实谱块分类和维数语义，QR/SVD 解决相位传播下的子空间迭代与收敛诊断。其四，在 Halo N45、Vertical N57、Sun–Earth 468 等案例上两种改进方法有效，并在独立 MATLAB 后端与全新进程重跑中复核。其五，physical Route H 四案和低分辨率全片收敛仍失败，失败未被隐藏。

其六，失败来源需要分层：点式方法主要是方程语义错误；Route H 一维失败主要受真实二维复谱结构约束，并伴随部分源状态边界；legacy 68 还存在初始化敏感；低分辨率流形片差异则是全局几何/分辨率问题。其七，局部 bundle 收敛只控制切空间种子，非线性传播会积累曲率与采样误差，因此全局 manifold sheet 必须使用独立几何门槛。

在冻结阈值下，partial real-Schur 和 shifted QR/SVD 相对 pointwise eig 显著降低了可接受案例的 cocycle 残差，同时把二维共轭子空间、初始化敏感和全片不收敛等负结果明确暴露出来。最稳妥的论文定位是 `numerical_framework_and_systematic_comparison`：贡献在于同源数据、统一门槛、独立后端、失败证据和可重复工程的系统结合，而非新理论。下一步应优先增加稳定子束、二维 Route H 流形、长事件传播和更多 Sun–Earth 独立案例，再由导师判断目标期刊与理论深化方向。

<!-- PAGEBREAK -->

## 参考文献

{reference_list(literature_config)}
"""


def claim_rows(metrics: dict[str, Any]) -> list[dict[str, str]]:
    return [
        {
            "claim_id": "C01",
            "claim_text": "54 图实现完整工程覆盖，但不是整篇 McCarthy 2018 严格等价复现。",
            "supporting_cases": "54 reproduction targets",
            "supporting_csv": "data/computed/reproduction_baseline_v1_summary.csv",
            "supporting_figure": "",
            "acceptance_threshold": "frozen evidence labels retained",
            "status": "supported_with_boundary",
            "limitation": "7 accepted; 30 boundary; 5 diagnostic; 12 proxy",
        },
        {
            "claim_id": "C02",
            "claim_text": "Benchmark 含 15 案例、4 类轨道族和 3 种方法。",
            "supporting_cases": "all registry cases",
            "supporting_csv": "research/invariant_bundles/benchmarks/benchmark_registry.csv;research/invariant_bundles/results/csv/method_comparison.csv",
            "supporting_figure": "figures/fig_bundle_method_summary.png",
            "acceptance_threshold": "15 registry rows; 45 method rows",
            "status": "supported",
            "limitation": "有限案例集，不是轨道族穷举",
        },
        {
            "claim_id": "C03",
            "claim_text": "Pointwise eig 为 0/15 通过；Schur 为 7 通过、4 边界、4 失败；QR/SVD 为 10 通过、5 失败。",
            "supporting_cases": "all 15 cases",
            "supporting_csv": "research/invariant_bundles/results/csv/method_comparison.csv",
            "supporting_figure": "figures/fig_bundle_method_summary.png",
            "acceptance_threshold": "max residual <=1e-6 accepted; <=1e-3 boundary absent other failure",
            "status": "supported",
            "limitation": "研究层状态不能提升冻结复现等级",
        },
        {
            "claim_id": "C04",
            "claim_text": "Halo N45、Vertical N57、Sun–Earth 468 上两种改进方法均通过，并显著降低 cocycle 残差。",
            "supporting_cases": "em_halo_12p40_n45;em_vertical_12p66_n57;se_active_geometry_member_468",
            "supporting_csv": "research/invariant_bundles/results/csv/method_comparison.csv",
            "supporting_figure": "figures/fig_phase_continuity_profiles.png",
            "acceptance_threshold": "max residual <=1e-6",
            "status": "supported",
            "limitation": "仅三个高分辨率锚点",
        },
        {
            "claim_id": "C05",
            "claim_text": "MATLAB R2024a 独立 real-Schur 后端在 12 个关键案例上维数、分类和状态全部一致。",
            "supporting_cases": "12 independent-backend cases",
            "supporting_csv": "research/invariant_bundles/results/csv/independent_schur_backend_comparison.csv",
            "supporting_figure": "",
            "acceptance_threshold": "principal angle <=1e-4 deg and exact categorical agreement",
            "status": "supported",
            "limitation": f"最大主角 {metrics['schur_max_angle_deg']:.3e} deg；有限离散算子验证",
        },
        {
            "claim_id": "C06",
            "claim_text": "四个 physical Route H QR/SVD 失败分类为 no_accepted_1d_bundle；legacy 68 为 method_initialization_sensitive。",
            "supporting_cases": "route_h_member_17;route_h_member_32;route_h_member_54;route_h_member_68;route_h_member_68_legacy_dg_positive",
            "supporting_csv": "research/invariant_bundles/results/csv/qr_svd_failure_classification.csv",
            "supporting_figure": "figures/fig_route_h_rho_control.png",
            "acceptance_threshold": "3 initializations; caps 200/500/1000; native N45 and diagnostic N67; 80-digit residual recomputation",
            "status": "supported_with_bounded_scope",
            "limitation": "不是所有算法/分辨率上的不可能性证明",
        },
        {
            "claim_id": "C07",
            "claim_text": "仅符号/相位对齐不能修复点式 eig 的 cocycle 语义错误。",
            "supporting_cases": "5 ablation cases",
            "supporting_csv": "research/invariant_bundles/results/csv/ablation_study.csv",
            "supporting_figure": "figures/ablation_bundle_residual.png;figures/ablation_phase_continuity.png",
            "acceptance_threshold": "frozen residual thresholds",
            "status": "supported",
            "limitation": "消融限于 5 案例×7 变体",
        },
        {
            "claim_id": "C08",
            "claim_text": "Schur 相位跟踪改善 Route H 二维子空间表示，但没有产生可接受的一维 bundle。",
            "supporting_cases": "route_h_member_32;route_h_member_68",
            "supporting_csv": "research/invariant_bundles/results/csv/ablation_study.csv",
            "supporting_figure": "figures/ablation_bundle_residual.png",
            "acceptance_threshold": "dimension semantics retained; max residual <=1e-6 for acceptance",
            "status": "supported",
            "limitation": "改善不等于通过",
        },
        {
            "claim_id": "C09",
            "claim_text": "Stage F 保存 126 行流形结果，其中 36 行通过、90 行失败，失败行未隐藏。",
            "supporting_cases": "7 manifold cases",
            "supporting_csv": "research/invariant_bundles/results/csv/manifold_convergence.csv",
            "supporting_figure": "figures/fig_manifold_method_metrics.png",
            "acceptance_threshold": "source, bundle, integration, geometry gates all satisfied",
            "status": "supported",
            "limitation": "一维不稳定子束和固定传播窗",
        },
        {
            "claim_id": "C10",
            "claim_text": "局部 bundle 收敛不保证低分辨率全局 manifold sheet 通过 0.01 几何门槛。",
            "supporting_cases": "Halo N21/N33/N45;Vertical N33/N45/N57",
            "supporting_csv": "research/invariant_bundles/results/csv/manifold_convergence.csv",
            "supporting_figure": "figures/fig_halo_manifold_displacement_sheets.png",
            "acceptance_threshold": "cross-resolution normalized 3D distance <=0.01",
            "status": "supported",
            "limitation": "有限传播时间和采样网格",
        },
        {
            "claim_id": "C11",
            "claim_text": "全新进程重跑的 5301 项科学检查全部通过且最大科学数值相对差为 0。",
            "supporting_cases": "15 bundle cases;7 manifold cases",
            "supporting_csv": "research/invariant_bundles/independent_rerun/comparison_to_stage_f.csv",
            "supporting_figure": "",
            "acceptance_threshold": "numeric atol=1e-12 rtol=1e-8; exact categorical agreement",
            "status": "supported",
            "limitation": "实现可重复性不等于论文正确性或投稿条件",
        },
        {
            "claim_id": "C12",
            "claim_text": "Physical corrected-rho Route H 是二维实共轭子空间/失败；legacy seed-rho 是一维/通过正控制。",
            "supporting_cases": "route_h physical 17/32/54/68;legacy 68",
            "supporting_csv": "research/invariant_bundles/results/csv/independent_schur_backend_comparison.csv;research/invariant_bundles/results/csv/method_comparison.csv",
            "supporting_figure": "figures/fig_route_h_rho_control.png",
            "acceptance_threshold": "independent dimension and status agreement",
            "status": "supported",
            "limitation": "legacy 控制不能替代物理算子",
        },
        {
            "claim_id": "C13",
            "claim_text": "Chapter 4 projection holdout 保持 0/4，paper_projection=fail，paper_3d=false。",
            "supporting_cases": "Fig.4.3-4.6 frozen holdout panels",
            "supporting_csv": "data/computed/chapter4_fig43_fig46_projection_holdout_audit.csv",
            "supporting_figure": "",
            "acceptance_threshold": "all frozen camera/epsilon/crop thresholds",
            "status": "supported_negative",
            "limitation": "任何研究后验结果都不得覆盖",
        },
        {
            "claim_id": "C14",
            "claim_text": "论文定位为数值框架与系统比较，而非方法学首创。",
            "supporting_cases": "25 verified formal references;all experimental stages",
            "supporting_csv": "research/invariant_bundles/paper/literature_matrix.csv",
            "supporting_figure": "",
            "acceptance_threshold": "9 required literature topics covered; no fabricated DOI",
            "status": "supported",
            "limitation": "文献矩阵不是穷尽性系统综述",
        },
        {
            "claim_id": "C15",
            "claim_text": "当前材料可供导师完整评审，但不声明已经达到投稿条件。",
            "supporting_cases": "all",
            "supporting_csv": "research/invariant_bundles/paper_release/claim_evidence_matrix.csv",
            "supporting_figure": "",
            "acceptance_threshold": "explicit limitations and unresolved questions retained",
            "status": "supported_with_boundary",
            "limitation": "需补稳定子束、二维流形、长事件传播和期刊定位",
        },
    ]


def limitations_markdown(config: dict[str, Any], metrics: dict[str, Any]) -> str:
    return f"""# 局限性与真实性边界

1. **复现等级被冻结。** 54 图为完整工程覆盖，不是整篇 McCarthy 2018 严格数值等价。证据仍为 7 accepted、30 boundary、5 diagnostic、12 proxy，研究实验不得回写提升。
2. **Chapter 4 不放宽。** projection holdout 保持 `0/4`、`paper_projection=fail`、`paper_3d=false`；后验 12.397983 日 N21 候选仅是诊断线索。
3. **没有新理论。** 本文没有证明 bundle 的存在、唯一性、可约化性或收敛率；定位固定为 `{config['study_positioning']}`。
4. **独立后端范围有限。** MATLAB R2024a `schur/ordschur` 在 12 个离散算子上 12/12 一致，最大主角 {metrics['schur_max_angle_deg']:.3e}°；这不是连续问题定理，也不是第三种硬件/库后端。
5. **QR/SVD 失败分类是有界结论。** 测试了三种初始化、200/500/1000 次迭代、N45 与诊断 N67；80 位运算只重算残差，没有任意精度重积分轨迹。
6. **Route H 负结果保留。** corrected-rho physical 17/32/54/68 是二维实共轭子空间并保持一维失败；legacy seed-rho 只作为正控制，不能替代物理算子。
7. **流形范围有限。** 仅验证 7 案例的一维不稳定子束、固定传播窗、三种扰动和两个符号；稳定子束、二维流形对象、事件终止和长期传播尚缺。
8. **局部与全局不等价。** Halo/Vertical 低分辨率全片距离仍超过 0.01；局部 bundle 通过不能覆盖全局 manifold sheet 失败。
9. **文献不是穷尽性综述。** 25 项正式来源覆盖九主题，21 个 DOI 核实、4 个 DOI 明确未分配；不能据此声称方法首创。
10. **可重复不等于可投稿。** 全新进程 5301 项科学比较一致、CI 守护权威文件，但仍需理论深化、稳定子束、二维 Route H 流形、更多 Sun–Earth 案例和目标期刊评估。

## 明确停止边界

本阶段只生成可交导师评审的中文初稿。独立验证虽已完成，但本稿仍标记 `submission_readiness=not_claimed`，不得在导师审阅和缺口补齐前改写为“已经达到投稿条件”。
"""


def reviewer_markdown(config: dict[str, Any], metrics: dict[str, Any]) -> str:
    return f"""# 导师/审稿人快速评估

## 一句话判断

材料已形成完整、可追溯的中文论文初稿，可用于导师判断研究定位和下一轮实验；当前不宜声称投稿就绪。

## 已经站得住的部分

- 15 案例×3 方法的结果完整，包含 {metrics['status_counts']['ordered_partial_real_schur_tracking']['accepted']} 个 Schur 通过、{metrics['status_counts']['qr_svd_shifted_cocycle_iteration']['accepted']} 个 QR/SVD 通过和全部负结果。
- MATLAB R2024a 独立 Schur 后端在 12 个关键案例上分类、维数和状态一致。
- 五个 QR/SVD 失败案例经过初始化、迭代上限、分辨率和 80 位残差重算分类。
- 35 行消融、126 行流形传播、全新进程重跑和 CI 均有 CSV/NPZ/日志/哈希。
- 25 项正式文献覆盖九主题，论文定位没有包装成方法首创。

## 需要导师重点判断的部分

1. “数值框架与系统比较”是否足以作为主贡献，还是应把“失败模式与算子语义诊断”提升为副标题。
2. Route H corrected-rho 的二维共轭子空间是否应成为正文核心案例，还是放入专门讨论/附录。
3. 下一轮优先补稳定一维 bundle，还是直接实现二维 Route H manifold sheet。
4. 是否扩展 Sun–Earth 案例并形成独立应用章节。
5. 在补实验前，目标应优先中文核心、国内会议，还是继续按英文期刊结构深化。

## 不能被评审措辞覆盖的事实

- McCarthy 复现等级被冻结；54 图不是整篇严格等价。
- Chapter 4 仍是 `0/4`、`paper_projection=fail`、`paper_3d=false`。
- Physical Route H 四案仍为二维/失败；legacy 68 仅是正控制。
- 本稿状态为 `not_submission_ready`。
"""


def clean_inline(text: str) -> str:
    text = text.replace("`", "").replace("**", "").replace("*", "")
    return text.replace("\\|", "|")


def set_run_east_asia_font(run: Any, font_name: str) -> None:
    """Set an East Asian run font even when python-docx has no rFonts node yet."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    run.font.name = font_name
    run_properties = run._element.get_or_add_rPr()
    fonts = run_properties.rFonts
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        run_properties.insert(0, fonts)
    fonts.set(qn("w:eastAsia"), font_name)


def set_cell_text(cell: Any, text: str, *, bold: bool = False, size: float = 8.5) -> None:
    from docx.shared import Pt

    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(clean_inline(text))
    run.bold = bold
    run.font.size = Pt(size)
    set_run_east_asia_font(run, "宋体")


def configure_docx(document: Any, config: dict[str, Any]) -> None:
    from docx.enum.section import WD_SECTION
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt

    section = document.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.6)
    section.right_margin = Cm(2.4)
    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(10.5)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.first_line_indent = Cm(0.74)
    normal.paragraph_format.space_after = Pt(0)
    for style_name, east_asia, size in (
        ("Title", "黑体", 22),
        ("Heading 1", "黑体", 16),
        ("Heading 2", "黑体", 14),
        ("Heading 3", "黑体", 12),
    ):
        style = document.styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
        style.paragraph_format.keep_with_next = True
    core = document.core_properties
    core.title = config["title"]
    core.author = config["author"]
    core.subject = "拟周期轨道实不变子束数值比较与可靠性分析"
    core.keywords = "拟周期轨道; 不变子束; Schur; QR/SVD; CR3BP"
    for current_section in document.sections:
        footer = current_section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer.add_run()
        fld_char1 = OxmlElement("w:fldChar")
        fld_char1.set(qn("w:fldCharType"), "begin")
        instr_text = OxmlElement("w:instrText")
        instr_text.set(qn("xml:space"), "preserve")
        instr_text.text = " PAGE "
        fld_char2 = OxmlElement("w:fldChar")
        fld_char2.set(qn("w:fldCharType"), "end")
        run._r.extend([fld_char1, instr_text, fld_char2])


def add_cover(document: Any, config: dict[str, Any]) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.shared import Pt

    for _ in range(5):
        document.add_paragraph()
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(config["title"])
    run.bold = True
    run.font.size = Pt(22)
    set_run_east_asia_font(run, "黑体")
    for _ in range(4):
        document.add_paragraph()
    for label, value in (
        ("作者", config["author"]),
        ("导师", config["adviser"]),
        ("单位", config["institution"]),
        ("日期", config["date"]),
    ):
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.first_line_indent = None
        run = paragraph.add_run(f"{label}：{value}")
        run.font.size = Pt(14)
        set_run_east_asia_font(run, "宋体")
    document.add_paragraph()
    status = document.add_paragraph()
    status.alignment = WD_ALIGN_PARAGRAPH.CENTER
    status.paragraph_format.first_line_indent = None
    run = status.add_run("导师评审初稿｜数值框架与系统比较｜not_submission_ready｜不声明投稿就绪")
    run.font.size = Pt(11)
    run.italic = True
    set_run_east_asia_font(run, "楷体")
    document.add_page_break()


def add_manual_contents(document: Any, config: dict[str, Any]) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    heading = document.add_heading("目录", level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for index, section in enumerate(config["required_sections"], start=1):
        p = document.add_paragraph()
        p.paragraph_format.first_line_indent = None
        p.paragraph_format.left_indent = None
        p.add_run(f"{index}. {section}").font.size = Pt(11)
    p = document.add_paragraph()
    p.paragraph_format.first_line_indent = None
    p.add_run("参考文献").font.size = Pt(11)
    document.add_page_break()


def add_markdown_paragraph(document: Any, text: str, *, style: str | None = None) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn

    paragraph = document.add_paragraph(style=style)
    if style in ("List Bullet", "List Number"):
        paragraph.paragraph_format.first_line_indent = None
    if text.startswith("表 ") or text.startswith("图 "):
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.first_line_indent = None
    parts = re.split(r"(\*\*.*?\*\*|`.*?`)", text)
    for part in parts:
        if not part:
            continue
        bold = part.startswith("**") and part.endswith("**")
        code = part.startswith("`") and part.endswith("`")
        value = part[2:-2] if bold else part[1:-1] if code else part
        run = paragraph.add_run(value)
        run.bold = bold
        if code:
            run.font.name = "Consolas"
        set_run_east_asia_font(run, "宋体")


def build_docx(markdown: str, target: Path, release: Path, config: dict[str, Any]) -> dict[str, int]:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt

    document = Document()
    configure_docx(document, config)
    add_cover(document, config)
    lines = markdown.splitlines()
    table_count = 0
    image_count = 0
    inserted_contents = False
    index = 0
    while index < len(lines):
        line = lines[index].strip()
        if line.startswith("# ") or line.startswith("- 作者：") or line.startswith("- 导师：") or line.startswith("- 单位：") or line.startswith("- 日期：") or line.startswith("- 稿件定位：") or line.startswith("- 状态："):
            index += 1
            continue
        if line == "<!-- PAGEBREAK -->":
            document.add_page_break()
            index += 1
            continue
        if line.startswith("## 1. 引言") and not inserted_contents:
            document.add_page_break()
            add_manual_contents(document, config)
            inserted_contents = True
        if line.startswith("### "):
            document.add_heading(clean_inline(line[4:]), level=2)
            index += 1
            continue
        if line.startswith("## "):
            document.add_heading(clean_inline(line[3:]), level=1)
            index += 1
            continue
        image_match = re.fullmatch(r"!\[(.*?)\]\((.*?)\)", line)
        if image_match:
            path = release / image_match.group(2)
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.first_line_indent = None
            run = paragraph.add_run()
            run.add_picture(str(path), width=Inches(6.15))
            image_count += 1
            caption = document.add_paragraph()
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption.paragraph_format.first_line_indent = None
            caption_run = caption.add_run(f"图 {image_count}  {image_match.group(1)}")
            caption_run.font.size = Pt(9.5)
            index += 1
            continue
        if line.startswith("|") and index + 1 < len(lines) and lines[index + 1].strip().startswith("|"):
            raw_rows: list[list[str]] = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                raw_rows.append([cell.strip() for cell in lines[index].strip().strip("|").split("|")])
                index += 1
            if len(raw_rows) >= 2 and all(set(cell) <= {"-", ":"} for cell in raw_rows[1]):
                table_data = [raw_rows[0], *raw_rows[2:]]
            else:
                table_data = raw_rows
            table = document.add_table(rows=len(table_data), cols=len(table_data[0]))
            table.style = "Table Grid"
            table.autofit = True
            for row_i, values in enumerate(table_data):
                for col_i, value in enumerate(values):
                    set_cell_text(table.cell(row_i, col_i), value, bold=row_i == 0)
            table_count += 1
            continue
        if line.startswith("- "):
            add_markdown_paragraph(document, line[2:], style="List Bullet")
            index += 1
            continue
        if re.match(r"^\d+\. ", line):
            add_markdown_paragraph(document, re.sub(r"^\d+\. ", "", line), style="List Number")
            index += 1
            continue
        if line.startswith("$$"):
            equation = line.strip("$")
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.first_line_indent = None
            run = paragraph.add_run(equation)
            run.font.name = "Cambria Math"
            index += 1
            continue
        if not line:
            index += 1
            continue
        paragraph_lines = [line]
        index += 1
        while index < len(lines):
            next_line = lines[index].strip()
            if not next_line or next_line.startswith(("#", "- ", "|", "![", "$$", "<!--")) or re.match(r"^\d+\. ", next_line):
                break
            paragraph_lines.append(next_line)
            index += 1
        add_markdown_paragraph(document, " ".join(paragraph_lines))
    target.parent.mkdir(parents=True, exist_ok=True)
    document.save(target)
    return {"docx_tables": table_count, "docx_images": image_count}


def validate_docx(path: Path, config: dict[str, Any]) -> dict[str, Any]:
    with zipfile.ZipFile(path) as archive:
        names = archive.namelist()
        xml = archive.read("word/document.xml").decode("utf-8")
        media = [name for name in names if name.startswith("word/media/")]
    text = re.sub(r"<[^>]+>", "", xml)
    required = [config["title"], config["author"], config["adviser"], config["institution"]]
    required.extend(config["required_sections"])
    missing = [marker for marker in required if marker not in text]
    if missing:
        raise ValueError(f"DOCX text missing required markers: {missing}")
    if len(media) < 9:
        raise ValueError(f"DOCX embeds only {len(media)} images; expected at least 9")
    return {
        "docx_bytes": path.stat().st_size,
        "docx_media_count": len(media),
        "docx_text_characters": len(text),
        "docx_required_marker_count": len(required),
    }


def render_validation_pdf(docx: Path, evidence: Path) -> dict[str, Any]:
    from pypdf import PdfReader
    import pypdfium2 as pdfium

    soffice = Path(r"C:\Program Files\LibreOffice\program\soffice.exe")
    if not soffice.is_file():
        raise FileNotFoundError("LibreOffice soffice.exe is required for DOCX render validation")
    render_dir = evidence / "rendered"
    render_dir.mkdir(parents=True, exist_ok=True)
    # Keep LibreOffice's volatile registry/cache outside committed evidence.
    # Its extension registry can create very deep Windows paths that are not
    # reliably removable from a nested repository directory.
    profile = Path(tempfile.mkdtemp(prefix="lo_profile_paper_"))
    command = [
        str(soffice),
        "--headless",
        f"-env:UserInstallation={profile.resolve().as_uri()}",
        "--convert-to",
        "pdf",
        "--outdir",
        str(render_dir),
        str(docx),
    ]
    completed = subprocess.run(
        command,
        cwd=ROOT,
        text=True,
        stdout=subprocess.PIPE,
        stderr=subprocess.STDOUT,
        timeout=180,
        check=False,
    )
    shutil.rmtree(profile, ignore_errors=True)
    if completed.returncode != 0:
        raise RuntimeError(f"LibreOffice render failed: {completed.stdout}")
    pdf = render_dir / f"{docx.stem}.pdf"
    if not pdf.is_file():
        raise FileNotFoundError(f"LibreOffice did not create {pdf}: {completed.stdout}")
    reader = PdfReader(str(pdf))
    page_count = len(reader.pages)
    if page_count < 12:
        raise ValueError(f"rendered Chinese paper has only {page_count} pages")
    extracted = "".join(page.extract_text() or "" for page in reader.pages)
    pdf_doc = pdfium.PdfDocument(str(pdf))
    preview_indices = sorted({0, page_count // 2, page_count - 1})
    previews = []
    for page_index in preview_indices:
        page = pdf_doc[page_index]
        bitmap = page.render(scale=1.5)
        image = bitmap.to_pil()
        target = render_dir / f"preview_page_{page_index + 1:03d}.png"
        image.save(target)
        previews.append(str(target.relative_to(ROOT)).replace("\\", "/"))
        bitmap.close()
        page.close()
    pdf_doc.close()
    return {
        "libreoffice_command": command,
        "libreoffice_return_code": completed.returncode,
        "libreoffice_output": completed.stdout.strip(),
        "pdf_path": str(pdf.relative_to(ROOT)).replace("\\", "/"),
        "pdf_bytes": pdf.stat().st_size,
        "pdf_pages": page_count,
        "pdf_extracted_characters": len(extracted),
        "preview_paths": previews,
    }


def write_npz(path: Path, data: dict[str, Any], metrics: dict[str, Any]) -> None:
    method = data["method"]
    np.savez_compressed(
        path,
        schema_version=np.array(["invariant_bundle_paper_release_validation_npz_v1"]),
        case_ids=np.array([row["case_id"] for row in data["registry"]]),
        method_ids=np.array(list(METHODS)),
        method_status=np.array([row["research_status"] for row in method]),
        method_max_residual=np.array([float(row["max_invariance_residual"]) for row in method]),
        method_bundle_dimension=np.array([int(row["bundle_dimension"]) for row in method]),
        independent_schur_max_angle_deg=np.array([metrics["schur_max_angle_deg"]]),
        manifold_status=np.array([row["status"] for row in data["manifold"]]),
        fresh_scientific_checks=np.array([metrics["fresh_scientific_checks"]], dtype=np.int64),
        chapter4_holdout=np.array([0, 4], dtype=np.int64),
        positioning=np.array(["numerical_framework_and_systematic_comparison"]),
        submission_readiness=np.array(["not_claimed"]),
    )


def write_hash_manifest(
    release: Path,
    evidence: Path,
    config_path: Path,
) -> None:
    target = evidence / "artifact_hashes.csv"
    inputs = [
        config_path,
        Path(__file__).resolve(),
        ROOT / "tests" / "test_invariant_bundle_paper_release.py",
        LITERATURE_CONFIG,
        LITERATURE_MATRIX,
        VERIFIED_BIB,
        REGISTRY,
        METHOD_CSV,
        MANIFOLD_CSV,
        RESOLUTION_CSV,
        RUNTIME_CSV,
        SCHUR_CSV,
        QR_FAILURE_CSV,
        ABLATION_CSV,
        FRESH_COMPARISON,
        REPRO_SUMMARY,
        CHAPTER4_HOLDOUT,
    ]
    outputs = [path for path in sorted(release.rglob("*")) if path.is_file()]
    outputs += [
        path
        for path in sorted(evidence.rglob("*"))
        if path.is_file() and path != target
    ]
    rows = [
        {
            "schema_version": "paper_release_stage_artifact_hash_v1",
            "path": str(path.relative_to(ROOT)).replace("\\", "/"),
            "bytes": path.stat().st_size,
            "sha256": sha256(path),
        }
        for path in inputs + outputs
    ]
    write_csv(target, rows)


def build(config_path: Path, release: Path, evidence: Path) -> dict[str, Any]:
    if release.exists() and any(release.iterdir()):
        raise RuntimeError(f"refusing to overwrite paper release: {release}")
    if evidence.exists() and any(evidence.iterdir()):
        raise RuntimeError(f"refusing to overwrite paper release evidence: {evidence}")
    release.mkdir(parents=True, exist_ok=True)
    evidence.mkdir(parents=True, exist_ok=True)
    logs = evidence / "logs"
    logs.mkdir()
    started = datetime.now(timezone.utc)
    config = json.loads(config_path.read_text(encoding="utf-8"))
    validate_config(config)
    data = load_data()
    if len(data["registry"]) != 15 or len(data["method"]) != 45:
        raise ValueError("authoritative benchmark cardinality changed")
    if len(data["manifold"]) != 126 or len(data["schur"]) != 12:
        raise ValueError("authoritative validation cardinality changed")
    if len(data["qr_failure"]) != 5 or len(data["ablation"]) != 35:
        raise ValueError("failure-classification or ablation cardinality changed")
    if len(data["holdout"]) != 4 or any(
        row["paper_projection_acceptance"] != "fail"
        or row["paper_3d_equivalence"] != "false"
        for row in data["holdout"]
    ):
        raise ValueError("frozen Chapter 4 holdout truth boundary changed")
    metrics = compute_metrics(data)
    tables = build_summary_tables(release, data, metrics)
    figures = copy_figures(config, release)
    markdown = manuscript_markdown(config, data, metrics)
    manuscript = release / "manuscript_zh.md"
    manuscript.write_text(markdown, encoding="utf-8")
    claims = claim_rows(metrics)
    write_csv(release / "claim_evidence_matrix.csv", claims, CLAIM_FIELDS)
    (release / "limitations.md").write_text(
        limitations_markdown(config, metrics), encoding="utf-8"
    )
    (release / "reviewer_quick_assessment.md").write_text(
        reviewer_markdown(config, metrics), encoding="utf-8"
    )
    shutil.copy2(VERIFIED_BIB, release / "references.bib")
    docx = release / "manuscript_zh.docx"
    docx_metrics = build_docx(markdown, docx, release, config)
    docx_metrics.update(validate_docx(docx, config))
    render_metrics = render_validation_pdf(docx, evidence)
    (evidence / "render_validation.json").write_text(
        json.dumps(render_metrics, indent=2, ensure_ascii=False, sort_keys=True) + "\n",
        encoding="utf-8",
    )
    validation_npz = evidence / "paper_release_validation.npz"
    write_npz(validation_npz, data, metrics)
    checks = [
        {"check_id": "identity_fields", "status": "pass", "detail": "兀文昊;张晨;中国科学院大学"},
        {"check_id": "required_sections", "status": "pass", "detail": "13/13"},
        {"check_id": "formal_references", "status": "pass", "detail": f"{metrics['reference_count']} references; {metrics['doi_count']} verified DOI"},
        {"check_id": "claim_evidence_rows", "status": "pass", "detail": str(len(claims))},
        {"check_id": "figures", "status": "pass", "detail": f"{len(figures)} raster plus vector sources"},
        {"check_id": "summary_tables", "status": "pass", "detail": str(len(tables))},
        {"check_id": "docx_embedded_images", "status": "pass", "detail": str(docx_metrics['docx_media_count'])},
        {"check_id": "docx_render_pages", "status": "pass", "detail": str(render_metrics['pdf_pages'])},
        {"check_id": "chapter4_holdout", "status": "pass", "detail": "0/4;paper_projection=fail;paper_3d=false"},
        {"check_id": "route_h_boundary", "status": "pass", "detail": "physical=2D/fail;legacy=1D/accepted control"},
        {"check_id": "submission_readiness", "status": "pass", "detail": "not_claimed"},
    ]
    write_csv(evidence / "paper_release_checks.csv", checks)
    (evidence / "failure_evidence.md").write_text(
        "# Chinese paper release failure and boundary evidence\n\n"
        "## Retained negative results\n\n"
        "- Pointwise eig remains 0/15 accepted.\n"
        "- Four physical corrected-rho Route H cases remain two-dimensional/fail.\n"
        "- The legacy seed-rho case remains a one-dimensional accepted positive control and is not substituted for the physical operator.\n"
        "- QR/SVD failure classification retains four `no_accepted_1d_bundle` rows and one `method_initialization_sensitive` row.\n"
        "- Manifold evidence retains 90 failed rows; lower-resolution full sheets remain above 0.01.\n"
        "- Chapter 4 remains 0/4 with `paper_projection=fail` and `paper_3d=false`.\n\n"
        "## Scope statement\n\n"
        "The manuscript is complete for adviser review but explicitly records `submission_readiness=not_claimed`. Independent validation does not erase the listed scientific limitations.\n",
        encoding="utf-8",
    )
    environment = {
        "schema_version": "paper_release_stage_environment_v1",
        "python": sys.version,
        "python_executable": sys.executable,
        "platform": platform.platform(),
        "numpy": np.__version__,
        "libreoffice": r"C:\Program Files\LibreOffice\program\soffice.exe",
        "document_builder": "python-docx",
        "pdf_renderer": "LibreOffice headless",
    }
    try:
        import docx as docx_module

        environment["python_docx"] = docx_module.__version__
    except Exception as exc:  # pragma: no cover - build cannot reach here without docx
        environment["python_docx"] = f"unavailable: {exc}"
    (logs / "environment.json").write_text(
        json.dumps(environment, indent=2, ensure_ascii=False, sort_keys=True) + "\n",
        encoding="utf-8",
    )
    finished = datetime.now(timezone.utc)
    summary = {
        "schema_version": "invariant_bundle_paper_release_summary_v1",
        "status": "pass",
        "title": config["title"],
        "author": config["author"],
        "adviser": config["adviser"],
        "institution": config["institution"],
        "positioning": config["study_positioning"],
        "submission_readiness": "not_claimed",
        "required_sections": 13,
        "claim_evidence_rows": len(claims),
        "formal_references": metrics["reference_count"],
        "verified_doi": metrics["doi_count"],
        "release_figures": len(figures),
        "summary_tables": len(tables),
        "docx_pages_rendered": render_metrics["pdf_pages"],
        "docx_images": docx_metrics["docx_media_count"],
        "chapter4_holdout": "0/4",
        "truth_boundary_status": "preserved",
    }
    (evidence / "paper_release_summary.json").write_text(
        json.dumps(summary, indent=2, ensure_ascii=False, sort_keys=True) + "\n",
        encoding="utf-8",
    )
    (logs / "stage_execution.log").write_text(
        "paper release build\n"
        f"started_utc={started.isoformat()}\n"
        f"finished_utc={finished.isoformat()}\n"
        f"elapsed_seconds={(finished-started).total_seconds():.6f}\n"
        f"python={sys.executable}\n"
        f"title={config['title']}\n"
        f"identity={config['author']}|{config['adviser']}|{config['institution']}\n"
        "sections=13/13\n"
        f"claims={len(claims)}\n"
        f"references={metrics['reference_count']}\n"
        f"figures={len(figures)}\n"
        f"tables={len(tables)}\n"
        f"docx_pages={render_metrics['pdf_pages']}\n"
        "chapter4_holdout=0/4\n"
        "truth_boundary_status=preserved\n"
        "submission_readiness=not_claimed\n",
        encoding="utf-8",
    )
    write_hash_manifest(release, evidence, config_path)
    return summary


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--config", type=Path, default=DEFAULT_CONFIG)
    parser.add_argument("--release-dir", type=Path, default=DEFAULT_RELEASE)
    parser.add_argument("--evidence-dir", type=Path, default=DEFAULT_EVIDENCE)
    parser.add_argument("--check-only", action="store_true")
    args = parser.parse_args()
    config = json.loads(args.config.resolve().read_text(encoding="utf-8"))
    validate_config(config)
    if args.check_only:
        print(json.dumps({"status": "pass", "sections": 13, "figures": len(config["figures"])}, indent=2))
        return 0
    summary = build(args.config.resolve(), args.release_dir.resolve(), args.evidence_dir.resolve())
    print(json.dumps(summary, indent=2, ensure_ascii=False, sort_keys=True))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
